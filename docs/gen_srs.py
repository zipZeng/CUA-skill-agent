"""
Generate 需求规格说明书 (SRS) for CUA-Skill Agent project.
Target: 20+ pages, .docx format.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# --- Style helpers ---
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        run.font.color.rgb = RGBColor(255,255,255)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph('')
    return table

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('需求规格说明书')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CUA-Skill：面向桌面环境的计算机使用Agent技能库')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')
doc.add_paragraph('')

info_lines = [
    f'版本：V1.0',
    f'日期：{datetime.date.today().strftime("%Y-%m-%d")}',
    '状态：初稿',
    '密级：内部',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# REVISION HISTORY
# ============================================================
add_heading_styled(doc, '修订记录', 1)
add_table(doc,
    ['版本', '日期', '修订章节', '修订内容', '修订人'],
    [
        ['V1.0', datetime.date.today().strftime('%Y-%m-%d'), '全部', '初稿编写', '项目组'],
    ]
)
doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
add_heading_styled(doc, '目录', 1)
add_para(doc, '（目录由 Word 自动生成，请在 Word 中右键此处 → 更新域）')
doc.add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_heading_styled(doc, '第一章 引言', 1)

add_heading_styled(doc, '1.1 编写目的', 2)
add_para(doc, '本文档旨在对"CUA-Skill：面向桌面环境的计算机使用Agent技能库"系统进行全面的需求分析，明确系统的功能需求、非功能需求、数据需求和外部接口需求。本文档作为项目设计、开发、测试和验收的依据，指导项目团队在后续阶段的工作。', indent=True)
add_para(doc, '本文档的预期读者包括：', indent=True)
for reader in ['软件开发工程师：理解系统需要实现的具体功能和行为；', '测试工程师：根据需求设计测试用例和验证方案；', '架构师：在需求约束下进行技术选型和架构设计；', '项目管理人员：评估项目范围和工作量；', '学术评审人员：了解项目的技术深度和完成度。']:
    add_para(doc, f'  ● {reader}')

add_heading_styled(doc, '1.2 项目背景', 2)
add_para(doc, '当前AI正经历从文本生成到直接系统交互的范式转变。计算机使用Agent（Computer-Using Agent, CUA）旨在通过图形用户界面（GUI）自主操作计算机，完成文档编辑、网页导航、数据分析与系统配置等真实桌面任务。据行业监测，2026年春，数十万名用户排队等待安装桌面AI代理工具，这场变革背后，真正的商业价值不在聊天框里，而在桌面自动化的自主执行中。', indent=True)
add_para(doc, '然而，现有Agent系统难以规模化，关键瓶颈在于缺乏可复用的、结构化的技能抽象来捕捉人类如何使用图形用户界面。面对几十个相互依赖的动作序列，微小偏差就可能连锁式失败。针对这一痛点，微软研究院于2026年1月开源了CUA-Skill——一个大规模的技能库，编码了覆盖常用Windows应用程序的人类计算机使用知识，并在WindowsAgentArena基准上达到了57.5%的成功率。', indent=True)
add_para(doc, '2026年，Agent生态在桌面领域迎来多项突破性进展：Microsoft CUA-Skill构建了首个大规模桌面应用技能库；OpenClaw在GitHub上获得超过10万星标；腾讯云首次发布涵盖基础设施到应用的Agent产品全景图；MCP标准化工具协议得到普遍采纳，ZPIT-desktop-MCP等高性能桌面控制服务以Rust实现核心控制能力。', indent=True)
add_para(doc, '本项目的核心价值在于：直面2026年计算机使用Agent的前沿课题，构建一个真正具备可复用、可扩展技能库的桌面自动化Agent系统，从技能编码→Agent推理→执行验证，完整实现Agentic桌面交互的全链条。', indent=True)

add_heading_styled(doc, '1.3 术语与缩写', 2)
add_table(doc,
    ['术语/缩写', '全称', '说明'],
    [
        ['CUA', 'Computer-Using Agent', '计算机使用Agent，能够通过GUI自主操作计算机的AI系统'],
        ['Skill Node', '技能节点', '将人类桌面操作知识编码为可参数化、可组合执行的最小功能单元'],
        ['Skill Base', '技能库', '存储和组织所有技能节点的仓库，支持语义检索和动态加载'],
        ['MCP', 'Model Context Protocol', '模型上下文协议，2026年AI Agent工具调用的核心通信标准'],
        ['UIA', 'UI Automation', 'Windows辅助功能API，用于获取桌面UI控件树结构'],
        ['LLM', 'Large Language Model', '大语言模型，用于任务分解和推理'],
        ['VLM', 'Vision-Language Model', '视觉语言模型，用于屏幕截图理解和元素识别'],
        ['OTel', 'OpenTelemetry', '开源可观测性框架，用于分布式链路追踪'],
        ['WAA', 'WindowsAgentArena', '标准化桌面Agent基准测试平台'],
    ]
)

add_heading_styled(doc, '1.4 参考资料', 2)
refs = [
    '[1] Microsoft Research, "CUA-Skill: A Skill Base for Computer-Using Agents", 2026.1',
    '[2] Anthropic, "Model Context Protocol (MCP) Specification", 2024-2026',
    '[3] ZPIT-desktop-MCP, "High-Performance Desktop Control MCP Server", GitHub, 2026',
    '[4] OpenDev Project, "Dual-Agent Architecture for Long-Horizon Tasks", 2026',
    '[5] WindowsAgentArena Benchmark Documentation, 2026',
    '[6] Ollama, "Local LLM Deployment Framework Documentation", 2026',
    '[7] OpenTelemetry, "Distributed Tracing Specification V1.0", 2026',
    '[8] Microsoft, "UI Automation Specification", Windows SDK Documentation',
]
for ref in refs:
    add_para(doc, ref)

doc.add_page_break()

# ============================================================
# CHAPTER 2: OVERALL DESCRIPTION
# ============================================================
add_heading_styled(doc, '第二章 总体描述', 1)

add_heading_styled(doc, '2.1 产品描述', 2)
add_para(doc, 'CUA-Skill Agent是一套面向桌面环境的计算机使用Agent系统，它将人类操作Windows/跨平台桌面应用的程序化知识编码为可复用、可参数化、可组合执行的技能库（Skill Base），覆盖文件管理、浏览器自动化、办公文档操作等常见应用场景。Agent支持动态技能检索、参数化实例化和内存感知故障恢复能力，能够理解和执行多步桌面任务。', indent=True)
add_para(doc, '系统核心特征：', indent=True)
for feat in [
    '技能驱动架构：将桌面交互知识结构化编码为Skill Node，支持语义检索和组合执行；',
    '分层执行框架：规划（Planner）与执行（Executor）分离，有效抑制上下文膨胀；',
    '混合感知策略：UI Automation（结构化元素定位）+ 视觉语言模型（非标准控件识别）双模感知；',
    '内存感知故障恢复：三级恢复策略（操作重试 → 技能替换 → 任务重规划），支持状态快照与回滚；',
    '标准化工具协议：基于MCP协议统一工具调用接口，集成ZPIT-desktop-MCP高性能桌面控制服务；',
    '全链路可观测：OpenTelemetry + Jaeger追踪Agent全执行链路，结构化日志支持审计与性能分析。',
]:
    add_para(doc, f'  ● {feat}')

add_heading_styled(doc, '2.2 用户特征', 2)
add_para(doc, '本系统的目标用户群体及特征如下：', indent=True)
add_table(doc,
    ['用户类别', '技术背景', '使用场景', '主要操作'],
    [
        ['终端用户', '具备基本计算机操作能力', '日常桌面任务自动化', '通过自然语言描述任务，查看执行结果'],
        ['开发者', '熟悉Python/Rust编程', '技能扩展与系统调试', '编写自定义Skill Node，调试执行链路，分析追踪日志'],
        ['系统管理员', '熟悉运维与监控', '批量桌面操作与审计', '批量任务编排，操作审计日志查询'],
        ['评估人员', '熟悉Agent评测方法', '基准测试与性能评估', '运行WindowsAgentArena评测，生成性能报告'],
    ]
)

add_heading_styled(doc, '2.3 运行环境', 2)
add_heading_styled(doc, '2.3.1 硬件环境', 3)
add_table(doc,
    ['组件', '最低配置', '推荐配置'],
    [
        ['CPU', 'Intel i5 8代 / AMD Ryzen 5', 'Intel i7 12代 / AMD Ryzen 7'],
        ['内存', '16 GB RAM', '32 GB RAM'],
        ['GPU', '无（使用云端视觉模型）', 'NVIDIA RTX 3060+ (8GB+ VRAM)，用于本地Ollama部署Qwen2-VL'],
        ['存储', '20 GB 可用空间', '50 GB SSD'],
        ['网络', '带宽 ≥ 10 Mbps', '带宽 ≥ 50 Mbps'],
    ]
)

add_heading_styled(doc, '2.3.2 软件环境', 3)
add_table(doc,
    ['软件项', '版本/说明'],
    [
        ['操作系统', 'Windows 10/11 (22H2+)'],
        ['Python', '3.10+'],
        ['Rust', '1.75+ (如需编译ZPIT-desktop-MCP)'],
        ['Ollama', '0.1.0+ (本地视觉模型运行时)'],
        ['Docker', '24.0+ (Jaeger、ChromaDB等服务编排)'],
        ['Git', '2.40+'],
    ]
)

add_heading_styled(doc, '2.4 设计与实现约束', 2)
add_table(doc,
    ['约束类别', '约束内容', '影响范围'],
    [
        ['技术栈约束', 'Agent运行时层必须使用Python 3.10+；桌面控制层优先使用Rust（可通过MCP Server解耦）', '架构设计'],
        ['协议约束', '工具调用必须遵循MCP标准协议；可观测性必须基于OpenTelemetry标准', '接口设计'],
        ['平台约束', '首期仅支持Windows 10/11；架构需预留macOS/Linux扩展点', '感知层、控制层'],
        ['模型约束', '支持Ollama本地部署（Qwen2-VL）和云端API（GPT-4o）两种视觉模型后端', '模型层'],
        ['性能约束', '单步操作延迟 < 5秒（不含视觉模型推理时间）；端到端任务成功率 ≥ 55%', '全部模块'],
        ['时间约束', '2周内完成核心功能开发和基准评测', '项目计划'],
        ['安全约束', '关键操作（文件删除、网络访问）需用户确认；所有操作结构化日志落盘审计', '执行引擎'],
    ]
)

add_heading_styled(doc, '2.5 假设与依赖', 2)
add_para(doc, '系统设计和开发基于以下假设：', indent=True)
assumptions = [
    'ZPIT-desktop-MCP服务可正常编译运行，提供截图、键鼠、窗口管理等基础桌面控制能力；',
    'Ollama可正常部署Qwen2-VL或类似视觉语言模型，推理延迟在可接受范围内（单帧<10秒）；',
    '目标桌面应用（文件管理器、Chrome/Edge浏览器、Microsoft Office）支持UI Automation接口；',
    'WindowsAgentArena基准测试环境可正常配置和运行；',
    '用户具备管理员权限，可安装必要的系统组件和驱动；',
    '网络环境可访问GitHub、PyPI等公共服务。',
]
for a in assumptions:
    add_para(doc, f'  ● {a}')

doc.add_page_break()

# ============================================================
# CHAPTER 3: FUNCTIONAL REQUIREMENTS (DETAILED)
# ============================================================
add_heading_styled(doc, '第三章 功能需求', 1)
add_para(doc, '本章按照系统架构分层，逐一详细描述各功能模块的需求。每个功能需求使用唯一编号标识，格式为 FR-XX-YY（XX=模块编号，YY=需求序号）。', indent=True)

# --- 3.1 Skill Base ---
add_heading_styled(doc, '3.1 技能库管理模块（FR-01）', 2)
add_para(doc, '技能库是本系统的核心资产，负责存储、组织、检索和管理所有桌面操作技能节点。', indent=True)

add_heading_styled(doc, '3.1.1 技能节点定义（FR-01-01）', 3)
add_para(doc, '系统应支持将每个桌面操作技能定义为结构化的SkillNode对象，包含以下字段：', indent=True)
add_table(doc,
    ['字段名', '类型', '必填', '说明'],
    [
        ['id', 'string', '是', '技能唯一标识符，采用snake_case命名，如"open_file_explorer"'],
        ['name', 'string', '是', '技能中文名称，如"打开文件管理器"'],
        ['description', 'string', '是', '技能功能描述，用于LLM语义理解和向量检索匹配'],
        ['category', 'enum', '是', '技能分类：file_ops | browser | office | system | custom'],
        ['parameters', 'list[Parameter]', '否', '参数定义列表，每个参数包含name/type/required/default/description'],
        ['preconditions', 'list[string]', '是', '执行前置条件列表，如["file_explorer_is_open"]'],
        ['postconditions', 'list[string]', '是', '执行后置条件列表，用于验证执行结果'],
        ['executable_ref', 'string', '是', '指向可执行代码的引用路径'],
        ['embedding', 'list[float]', '是', '技能的语义嵌入向量（维度=768或1024），用于向量检索'],
        ['timeout_ms', 'int', '否', '技能执行超时时间（毫秒），默认30000ms'],
        ['retry_policy', 'dict', '否', '重试策略配置：max_retries/backoff_factor/retry_on'],
        ['dependencies', 'list[string]', '否', '依赖的其他技能ID列表'],
    ]
)

add_heading_styled(doc, '3.1.2 技能定义格式（FR-01-02）', 3)
add_para(doc, '系统应支持以YAML格式声明式定义技能，允许非开发人员编写和维护技能。YAML定义文件应放置在skills-defs/目录下，按技能类别组织子目录。', indent=True)
add_para(doc, '每个YAML文件包含一个技能的完整定义，Agent运行时启动时自动加载并注册到技能注册表。', indent=True)

add_heading_styled(doc, '3.1.3 技能注册与管理（FR-01-03）', 3)
add_para(doc, '系统应提供技能注册表（SkillRegistry），支持以下操作：', indent=True)
add_table(doc,
    ['操作', '说明'],
    [
        ['register(skill)', '注册一个新的SkillNode到技能库'],
        ['unregister(skill_id)', '从技能库中移除指定技能'],
        ['get(skill_id)', '根据ID获取技能节点'],
        ['list_all()', '列出所有已注册的技能'],
        ['list_by_category(cat)', '按分类列出技能'],
        ['reload()', '重新从YAML文件加载所有技能定义'],
    ]
)

add_heading_styled(doc, '3.1.4 技能向量检索（FR-01-04）', 3)
add_para(doc, '系统应支持基于语义相似度的动态技能匹配。当用户输入自然语言任务指令时，系统将用户任务描述转化为嵌入向量，与技能库中所有技能的嵌入向量进行余弦相似度计算，返回Top-K最相关的技能节点。', indent=True)
add_para(doc, '具体要求：', indent=True)
for req in [
    '嵌入模型：使用sentence-transformers的all-MiniLM-L6-v2（维度384）或text2vec-base-chinese（中文优化）；',
    '向量存储：使用ChromaDB作为嵌入式向量数据库，支持持久化和增量更新；',
    '检索性能：Top-5检索延迟 < 100ms（技能库规模 < 100个技能）；',
    '相似度阈值：默认返回相似度 ≥ 0.6 的技能，低于阈值时提示用户补充描述；',
    '支持混合检索：结合语义相似度 + 分类过滤 + 参数匹配度的加权排序。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.1.5 预置技能集（FR-01-05）', 3)
add_para(doc, '系统应预置以下基础技能，覆盖文件管理和浏览器自动化两大核心场景：', indent=True)
add_table(doc,
    ['类别', '技能ID', '技能名称', '参数', '描述'],
    [
        ['文件管理', 'open_file_explorer', '打开文件管理器', 'path: string (optional)', '打开Windows资源管理器，可指定初始路径'],
        ['文件管理', 'list_files', '列出目录文件', 'path: string; pattern: string (optional)', '列出指定目录下的所有文件，支持通配符过滤'],
        ['文件管理', 'copy_file', '复制文件', 'source: string; dest: string', '将文件从源路径复制到目标路径'],
        ['文件管理', 'move_file', '移动文件', 'source: string; dest: string', '将文件从源路径移动到目标路径'],
        ['文件管理', 'rename_file', '重命名文件', 'path: string; new_name: string', '重命名指定文件或文件夹'],
        ['文件管理', 'create_folder', '创建文件夹', 'path: string; name: string', '在指定路径创建新文件夹'],
        ['文件管理', 'delete_file', '删除文件', 'path: string; permanent: bool', '删除文件（支持回收站或永久删除）'],
        ['文件管理', 'get_file_info', '获取文件属性', 'path: string', '获取文件大小、创建时间、修改时间等属性'],
        ['浏览器', 'open_browser', '打开浏览器', 'browser: string (optional)', '打开指定浏览器（默认系统默认浏览器）'],
        ['浏览器', 'navigate_to', '导航到URL', 'url: string', '在浏览器地址栏输入URL并导航'],
        ['浏览器', 'click_element', '点击元素', 'selector: string; by: enum', '根据选择器（xpath/css/text）点击页面元素'],
        ['浏览器', 'fill_form', '填写表单', 'fields: dict[{selector, value}]', '批量填写表单字段'],
        ['浏览器', 'get_page_content', '获取页面内容', 'selector: string (optional)', '获取页面文本内容或指定元素内容'],
        ['浏览器', 'scroll_page', '滚动页面', 'direction: enum; amount: int', '向指定方向滚动页面'],
        ['浏览器', 'take_screenshot', '截图', 'region: rect (optional)', '对当前浏览器窗口或指定区域截图'],
        ['浏览器', 'execute_js', '执行JS脚本', 'script: string', '在浏览器控制台执行JavaScript代码'],
    ]
)

add_heading_styled(doc, '3.1.6 技能组合图（FR-01-06）', 3)
add_para(doc, '系统应支持将多个基础技能组合为高层复合技能。组合方式包括：', indent=True)
for comp in [
    '顺序组合（Sequential）：按依赖关系依次执行各子技能，前一技能的输出可作为后一技能的输入参数；',
    '条件分支（Conditional）：根据执行结果或环境状态选择不同分支（if-else）；',
    '循环执行（Loop）：对多个对象重复执行同一技能（for-each）；',
    '并行执行（Parallel）：对无依赖关系的子技能并行执行，提升效率。',
]:
    add_para(doc, f'  ● {comp}')
add_para(doc, '组合图定义应支持JSON格式描述DAG（有向无环图）结构，包含节点（技能引用+参数绑定）和边（数据流/控制流）。', indent=True)

doc.add_page_break()

# --- 3.2 Agent Runtime ---
add_heading_styled(doc, '3.2 Agent核心运行时模块（FR-02）', 2)
add_para(doc, 'Agent核心运行时是整个系统的调度中枢，负责接收用户任务、检索匹配技能、实例化参数、执行技能序列、处理故障恢复和维护任务状态记忆。', indent=True)

add_heading_styled(doc, '3.2.1 任务规划器 Planner（FR-02-01）', 3)
add_para(doc, 'Planner负责将用户自然语言任务指令分解为可执行的技能序列（Skill Plan）。具体要求：', indent=True)
for req in [
    '接收用户自然语言任务指令（如"帮我整理桌面上最近3天创建的文件"）；',
    '调用LLM进行任务分解，输出结构化的步骤序列（Step List），每个步骤包含：skill_id / 参数绑定 / 预期前置条件 / 预期后置条件；',
    '支持多轮对话澄清模糊参数（如用户未指定目标文件夹时主动询问）；',
    '规划结果输出为JSON格式，包含任务ID、步骤列表、预估总耗时；',
    '规划延迟 < 5秒（不含LLM推理时间）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.2.2 参数解析器 Parameter Resolver（FR-02-02）', 3)
add_para(doc, 'Parameter Resolver负责根据当前UI上下文为技能参数动态填充值。具体要求：', indent=True)
for req in [
    '从Planner输出的参数绑定中提取参数名和值；',
    '若参数值缺失，尝试从UI上下文中推断（如当前活动窗口标题、选中的文件路径）；',
    '若无法从上下文推断，则提示用户补充；',
    '参数类型校验：在传递给技能执行前进行类型检查和格式验证；',
    '支持参数引用：一个步骤的输出可作为后续步骤的输入（如"步骤3使用步骤1创建的文件夹路径"）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.2.3 执行引擎 Execution Engine（FR-02-03）', 3)
add_para(doc, 'Execution Engine是技能执行的驱动核心，负责按照规划好的技能序列逐一调度执行。具体要求：', indent=True)
for req in [
    '支持四种执行模式：顺序执行、条件分支、循环执行、并行执行（独立子技能）；',
    '每个技能执行前进行前置条件检查，不满足时触发故障恢复；',
    '每个技能执行后进行后置条件验证，不满足时触发故障恢复；',
    '支持执行暂停/恢复/取消操作；',
    '记录每个技能执行的开始时间、结束时间、耗时、输入参数、输出结果、屏幕截图（执行前后各一张）；',
    '单技能执行超时控制：超过skill定义中的timeout_ms时强制中断并触发恢复流程。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.2.4 故障恢复 Failure Recovery（FR-02-04）', 3)
add_para(doc, '系统应实现三级故障恢复机制，确保在桌面操作失败时具有自愈能力：', indent=True)
add_table(doc,
    ['恢复级别', '触发条件', '恢复策略', '最大尝试次数'],
    [
        ['L1: 操作重试', 'UI元素未找到、网络超时、短暂状态不一致', '等待固定时间后重新执行当前技能操作', '3次'],
        ['L2: 技能替换', 'L1恢复失败、技能执行超时、工具调用不可用', '检索替代技能（语义相似 + 同类别）替换当前技能', '2次'],
        ['L3: 任务重规划', 'L2恢复失败、环境状态异常、用户取消', '通知Planner重新分解未完成子任务，必要时请求用户介入', '1次'],
    ]
)

add_heading_styled(doc, '3.2.5 记忆模块 Memory Module（FR-02-05）', 3)
add_para(doc, '记忆模块负责维护Agent的跨步骤和跨会话状态信息。具体要求：', indent=True)
for req in [
    '短期记忆（工作记忆）：存储当前任务的执行上下文，包括当前步骤索引、已完成步骤列表、环境状态快照（窗口布局、文件位置等）；',
    '长期记忆（持久化）：使用SQLite存储历史任务记录，包括任务描述、执行轨迹、成功/失败状态、耗时等；',
    '状态快照：在关键步骤前自动保存环境快照，支持故障时回滚到最近的稳定状态；',
    '跨会话恢复：支持从上次中断的任务继续执行（基于持久化的执行轨迹）。',
]:
    add_para(doc, f'  ● {req}')

doc.add_page_break()

# --- 3.3 GUI Perception ---
add_heading_styled(doc, '3.3 GUI感知模块（FR-03）', 2)
add_para(doc, 'GUI感知模块是Agent"看见"桌面环境的能力基础，提供UI元素定位和屏幕理解功能。', indent=True)

add_heading_styled(doc, '3.3.1 UI Automation接口（FR-03-01）', 3)
add_para(doc, '系统应封装Windows UI Automation API，提供结构化UI元素树获取能力。具体要求：', indent=True)
for req in [
    '获取当前活动窗口的完整UI元素树（Control Tree），包括元素类型、名称、位置、状态等属性；',
    '支持按条件查询UI元素：按Name、ClassName、AutomationId、ControlType筛选；',
    '支持获取元素的屏幕坐标位置（BoundingRectangle），用于鼠标点击定位；',
    '支持元素状态检测：是否可见（IsVisible）、是否可用（IsEnabled）、是否获得焦点（HasKeyboardFocus）；',
    '支持获取窗口列表：枚举所有顶级窗口，获取窗口标题、进程名、窗口状态。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.3.2 屏幕截图（FR-03-02）', 3)
add_para(doc, '系统应支持对桌面进行截图操作。具体要求：', indent=True)
for req in [
    '支持全屏截图和指定窗口/区域截图；',
    '截图格式：PNG（无损）或JPEG（压缩），分辨率与显示器一致；',
    '截图应包含鼠标光标位置指示；',
    '截图延迟 < 200ms（全屏，1920x1080分辨率）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.3.3 视觉模型理解（FR-03-03）', 3)
add_para(doc, '系统应集成视觉语言模型（VLM），支持基于屏幕截图的元素理解和定位。具体要求：', indent=True)
for req in [
    '支持多模态模型输入：将屏幕截图 + 自然语言查询（如"找到搜索按钮的位置"）发送给VLM；',
    '模型返回目标元素的屏幕坐标（bounding box）或描述性文本；',
    '支持Ollama本地部署（Qwen2-VL系列）和云端API（GPT-4o/GPT-4V）两种后端，通过配置切换；',
    '在UIA获取失败时自动回退到视觉模型定位（降级策略）；',
    '模型推理超时时间：本地模型 < 15秒/帧，云端模型 < 10秒/帧。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.3.4 混合感知策略（FR-03-04）', 3)
add_para(doc, '系统应采用混合感知策略，综合UIA和视觉模型的优势：', indent=True)
add_table(doc,
    ['场景', '感知方式', '优先级'],
    [
        ['标准Windows控件（按钮、文本框、菜单）', 'UI Automation', '1（首选）'],
        ['浏览器页面元素', 'UIA + WebDriver', '1（首选）'],
        ['Electron/Java等非标准控件', '视觉模型定位', '2（回退）'],
        ['图标、图片等非文本元素', '视觉模型定位', '2（回退）'],
        ['UIA不可用的应用程序', '视觉模型定位', '2（回退）'],
        ['模糊查询（"找到和XX相似的按钮"）', '视觉模型理解', '仅视觉'],
    ]
)

doc.add_page_break()

# --- 3.4 Desktop Control ---
add_heading_styled(doc, '3.4 桌面控制模块（FR-04）', 2)
add_para(doc, '桌面控制模块负责执行实际的桌面操作，包括键盘鼠标模拟、窗口管理和剪贴板操作。', indent=True)

add_heading_styled(doc, '3.4.1 键盘输入模拟（FR-04-01）', 3)
add_para(doc, '系统应支持以下键盘操作：', indent=True)
add_table(doc,
    ['操作', '说明', '参数'],
    [
        ['文本输入', '模拟键盘逐字符输入文本', 'text: string; interval_ms: int (字符间隔，默认50ms)'],
        ['按键组合', '模拟组合键按下（如Ctrl+C）', 'keys: list[string]; hold_ms: int (按住时间)'],
        ['特殊键', '模拟功能键（Enter/Tab/Esc等）', 'key: enum (SpecialKey)'],
        ['按键序列', '按顺序执行多个按键操作', 'sequence: list[KeyAction]'],
    ]
)

add_heading_styled(doc, '3.4.2 鼠标操作模拟（FR-04-02）', 3)
add_para(doc, '系统应支持以下鼠标操作：', indent=True)
add_table(doc,
    ['操作', '说明', '参数'],
    [
        ['鼠标移动', '将鼠标光标移动到指定坐标', 'x: int; y: int; duration_ms: int (移动耗时)'],
        ['鼠标点击', '在当前位置或指定坐标单击', 'button: enum (left/right/middle); x: int (optional); y: int (optional)'],
        ['鼠标双击', '在指定位置双击', 'x: int; y: int'],
        ['鼠标拖拽', '从起点拖拽到终点', 'start_x/y: int; end_x/y: int; duration_ms: int'],
        ['鼠标滚轮', '滚动鼠标滚轮', 'delta: int (正值向上); x: int (optional); y: int (optional)'],
    ]
)

add_heading_styled(doc, '3.4.3 窗口管理（FR-04-03）', 3)
add_para(doc, '系统应支持以下窗口管理操作：', indent=True)
add_table(doc,
    ['操作', '说明'],
    [
        ['枚举窗口', '获取当前所有顶级窗口列表（标题、句柄、位置、进程名）'],
        ['激活窗口', '将指定窗口切换到前台并获得焦点'],
        ['调整窗口', '移动窗口位置和/或调整大小'],
        ['最小化/最大化', '最小化、最大化或还原窗口'],
        ['关闭窗口', '关闭指定窗口（发送WM_CLOSE消息）'],
        ['获取前台窗口', '获取当前活动（前台）窗口的信息'],
    ]
)

add_heading_styled(doc, '3.4.4 剪贴板操作（FR-04-04）', 3)
add_para(doc, '系统应支持以下剪贴板操作：', indent=True)
for op in [
    '读取剪贴板文本内容；',
    '写入文本到剪贴板；',
    '读取剪贴板图片内容；',
    '清空剪贴板。',
]:
    add_para(doc, f'  ● {op}')

add_heading_styled(doc, '3.4.5 应用管理（FR-04-05）', 3)
add_para(doc, '系统应支持以下应用管理操作：', indent=True)
for op in [
    '启动应用程序（通过可执行文件路径或应用名）；',
    '通过进程名或窗口标题查找正在运行的应用；',
    '终止应用程序进程；',
    '等待应用窗口出现（超时控制）。',
]:
    add_para(doc, f'  ● {op}')

doc.add_page_break()

# --- 3.5 MCP Integration ---
add_heading_styled(doc, '3.5 MCP协议集成模块（FR-05）', 2)
add_para(doc, 'MCP（Model Context Protocol）是2026年AI Agent工具调用的核心通信标准，本系统需实现MCP客户端以调用ZPIT-desktop-MCP服务器提供的桌面控制工具。', indent=True)

add_heading_styled(doc, '3.5.1 MCP客户端（FR-05-01）', 3)
add_para(doc, '系统应实现MCP标准客户端，支持以下能力：', indent=True)
for req in [
    '使用stdio传输协议与MCP Server通信；',
    '支持标准的MCP握手流程（initialize → capabilities协商）；',
    '支持工具发现：通过tools/list获取Server提供的工具列表；',
    '支持工具调用：通过tools/call调用指定工具并获取返回值；',
    '连接健康检查：定期ping检测Server是否存活；',
    '自动重连：Server进程崩溃或重启时自动重连（最大重试3次，指数退避）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.5.2 MCP工具映射（FR-05-02）', 3)
add_para(doc, '系统应将MCP Server提供的工具映射为Agent内部统一的DesktopController接口。具体映射关系如下：', indent=True)
add_table(doc,
    ['DesktopController方法', 'MCP Tool名称', '说明'],
    [
        ['screenshot()', 'screenshot', '获取屏幕截图，返回base64编码图片'],
        ['mouse_click(x, y)', 'mouse_click', '在指定坐标执行鼠标左键单击'],
        ['mouse_double_click(x, y)', 'mouse_double_click', '在指定坐标执行鼠标左键双击'],
        ['mouse_move(x, y)', 'mouse_move', '将鼠标移动到指定坐标'],
        ['mouse_drag(x1,y1,x2,y2)', 'mouse_drag', '从(x1,y1)拖拽到(x2,y2)'],
        ['key_type(text)', 'key_type', '模拟键盘输入文本'],
        ['key_combo(keys)', 'key_combo', '执行组合键（如Ctrl+C）'],
        ['window_list()', 'window_list', '枚举所有窗口'],
        ['window_activate(hwnd)', 'window_activate', '激活指定窗口'],
        ['clipboard_get()', 'clipboard_get', '读取剪贴板文本'],
        ['clipboard_set(text)', 'clipboard_set', '设置剪贴板文本'],
        ['app_start(path)', 'app_start', '启动应用程序'],
        ['app_kill(name)', 'app_kill', '终止应用程序'],
    ]
)

doc.add_page_break()

# --- 3.6 Web Console ---
add_heading_styled(doc, '3.6 Web控制台模块（FR-06）', 2)
add_para(doc, 'Web控制台是用户与Agent交互的前端界面，提供任务提交、执行监控和历史查询功能。', indent=True)

add_heading_styled(doc, '3.6.1 任务提交界面（FR-06-01）', 3)
add_para(doc, '系统应提供Web界面供用户输入自然语言任务指令。要求：', indent=True)
for req in [
    '提供文本输入框，支持多行任务描述输入；',
    '提供示例任务快捷按钮（一键填充示例任务描述）；',
    '提交按钮触发后显示任务提交状态（成功/失败）；',
    '任务提交成功后自动跳转到执行监控页面。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.6.2 执行监控界面（FR-06-02）', 3)
add_para(doc, '系统应提供执行过程的实时可视化监控。要求：', indent=True)
for req in [
    '实时显示当前执行的技能步骤和整体进度（已完成/总数）；',
    '展示每步执行的前后截图对比；',
    '展示技能调用链路（时间线视图）；',
    '执行完成后展示总结（总耗时、成功率、关键截图）；',
    '支持执行过程的人工干预（暂停/继续/取消）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.6.3 历史记录查询（FR-06-03）', 3)
add_para(doc, '系统应提供历史任务记录的查询和回放功能。要求：', indent=True)
for req in [
    '支持按日期、任务状态、技能类别筛选历史记录；',
    '展示每个历史任务的执行摘要（任务描述、耗时、成功/失败状态）；',
    '支持查看历史任务的详细执行轨迹（每步的输入/输出/截图）；',
    '支持导出任务执行报告（JSON格式）。',
]:
    add_para(doc, f'  ● {req}')

doc.add_page_break()

# --- 3.7 Evaluation ---
add_heading_styled(doc, '3.7 评估与评测模块（FR-07）', 2)
add_para(doc, '评估模块负责对Agent系统进行标准化基准测试和性能分析。', indent=True)

add_heading_styled(doc, '3.7.1 WindowsAgentArena集成（FR-07-01）', 3)
add_para(doc, '系统应集成WindowsAgentArena基准测试框架，具体要求：', indent=True)
for req in [
    '能运行WindowsAgentArena提供的标准测试任务集；',
    '自动记录每个测试任务的成功/失败状态；',
    '统计关键指标：任务成功率、平均完成时间、平均操作步数、故障恢复率；',
    '生成评测报告（JSON + 图表）。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.7.2 自定义评测场景（FR-07-02）', 3)
add_para(doc, '除标准基准测试外，系统应支持自定义评测场景的定义和执行。自定义场景包括：', indent=True)
for scene in [
    '桌面文件整理：将桌面上符合特定条件的文件移动到指定文件夹；',
    '浏览器数据采集：打开指定网站，搜索关键词，提取并保存搜索结果；',
    '文档处理：打开指定文档，编辑内容，保存为不同格式。',
]:
    add_para(doc, f'  ● {scene}')
add_para(doc, '每个自定义场景需定义：任务描述、初始环境状态、期望最终状态、验证脚本。', indent=True)

add_heading_styled(doc, '3.7.3 性能指标定义（FR-07-03）', 3)
add_table(doc,
    ['指标名称', '计算方法', '目标值'],
    [
        ['任务成功率 (Success Rate)', '成功完成任务数 / 总任务数 × 100%', '≥ 55%'],
        ['平均完成时间 (Avg Time)', '所有任务完成时间之和 / 任务数', '< 3分钟/任务'],
        ['平均操作步数 (Avg Steps)', '所有任务执行步数之和 / 任务数', '< 15步/任务'],
        ['故障恢复率 (Recovery Rate)', '成功恢复的故障数 / 总故障数 × 100%', '> 70%'],
        ['单步执行延迟 (Step Latency)', '单步技能执行的平均耗时', '< 5秒'],
        ['技能检索延迟 (Retrieval Latency)', '技能检索的平均耗时', '< 100ms'],
    ]
)

doc.add_page_break()

# --- 3.8 Observability ---
add_heading_styled(doc, '3.8 可观测性模块（FR-08）', 2)
add_para(doc, '可观测性模块负责Agent执行链路的全量追踪和性能监控。', indent=True)

add_heading_styled(doc, '3.8.1 分布式链路追踪（FR-08-01）', 3)
add_para(doc, '系统应基于OpenTelemetry标准实现全链路追踪。要求：', indent=True)
for req in [
    '使用OpenTelemetry SDK在Python代码中自动或手动埋点；',
    '每个Span记录以下属性：span_name（技能名称/操作名称）、start_time、end_time、status（OK/ERROR）、attributes（参数、返回值等）；',
    'Span层级结构：Task Span > Step Span > Skill Span > Action Span（原子操作）；',
    '将Trace数据导出到Jaeger后端，通过Jaeger UI可视化展示调用链；',
    '支持采样策略：开发环境100%采样，生产环境可配置采样率。',
]:
    add_para(doc, f'  ● {req}')

add_heading_styled(doc, '3.8.2 结构化日志（FR-08-02）', 3)
add_para(doc, '系统应输出结构化日志（JSON格式），包含以下信息：', indent=True)
for log in [
    '时间戳（ISO 8601格式，精确到毫秒）；',
    '日志级别（DEBUG/INFO/WARNING/ERROR）；',
    'Trace ID和Span ID（关联到OpenTelemetry追踪）；',
    '模块名称（Planner/Executor/Retriever等）；',
    '消息内容（自由文本）；',
    '上下文数据（任务ID、技能ID、参数等）。',
]:
    add_para(doc, f'  ● {log}')

add_heading_styled(doc, '3.8.3 性能指标收集（FR-08-03）', 3)
add_para(doc, '系统应收集以下性能指标：', indent=True)
add_table(doc,
    ['指标类别', '指标名称', '类型', '说明'],
    [
        ['请求量', 'task_total', 'Counter', '接收到的任务总数'],
        ['请求量', 'task_success_total', 'Counter', '成功完成的任务数'],
        ['请求量', 'task_failure_total', 'Counter', '失败的任务数'],
        ['延迟', 'skill_execution_duration_ms', 'Histogram', '技能执行耗时分布'],
        ['延迟', 'retrieval_duration_ms', 'Histogram', '技能检索耗时分布'],
        ['延迟', 'vision_inference_duration_ms', 'Histogram', '视觉模型推理耗时分布'],
        ['错误', 'error_total', 'Counter', '各类错误计数（按error_type标签分类）'],
        ['资源', 'memory_usage_bytes', 'Gauge', 'Agent进程内存使用量'],
    ]
)

doc.add_page_break()

# ============================================================
# CHAPTER 4: NON-FUNCTIONAL REQUIREMENTS
# ============================================================
add_heading_styled(doc, '第四章 非功能需求', 1)

add_heading_styled(doc, '4.1 性能需求', 2)
add_table(doc,
    ['性能指标', '目标值', '测试条件'],
    [
        ['单步技能执行延迟', '< 5秒', '不含LLM/VLM推理时间'],
        ['屏幕截图延迟', '< 200ms', '1920×1080全屏截图'],
        ['技能检索延迟', '< 100ms', '技能库规模 < 100个技能'],
        ['任务规划延迟', '< 5秒', '不含LLM推理时间'],
        ['UIA元素定位延迟', '< 100ms', '单次元素查询'],
        ['视觉模型推理延迟（本地）', '< 15秒/帧', 'Qwen2-VL 7B, RTX 3060'],
        ['视觉模型推理延迟（云端）', '< 10秒/帧', 'GPT-4o API'],
        ['Web控制台响应延迟', '< 1秒', '页面加载和API响应'],
        ['Agent内存占用', '< 2GB', '正常运行状态'],
    ]
)

add_heading_styled(doc, '4.2 安全性需求', 2)
add_para(doc, '桌面操作Agent直接控制系统环境，安全性至关重要。系统应满足以下安全需求：', indent=True)
sec_reqs = [
    '操作确认机制：对于文件删除、系统配置修改、网络代理变更等高风险操作，必须在执行前弹出用户确认对话框，确认超时（30秒）后自动拒绝操作；',
    '操作白名单：支持配置允许执行的操作列表，白名单外的操作自动拒绝；',
    '路径沙箱：可配置允许访问的文件系统路径范围，超出范围的文件操作被拦截；',
    '操作审计：所有桌面操作（包括键盘输入、鼠标点击、文件修改）均记录结构化审计日志，包含时间戳、操作类型、操作目标、执行结果；',
    '敏感信息保护：审计日志中自动脱敏密码、API Key等敏感字段（正则匹配 + 字段标记）；',
    '指令注入防护：对用户输入的自然语言指令进行安全扫描，检测并拒绝包含危险操作（如格式化磁盘、修改注册表）的指令；',
    '权限最小化：Agent进程以最低必要权限运行，不请求管理员权限（除非特定操作必需）。',
]
for s in sec_reqs:
    add_para(doc, f'  ● {s}')

add_heading_styled(doc, '4.3 可用性需求', 2)
for ux in [
    '自然语言交互：用户无需学习特定命令语法，以自然语言描述任务即可；',
    '错误提示友好：当操作失败时，以自然语言向用户解释失败原因和可能的解决方案；',
    '进度可见：始终向用户展示当前的执行进度和状态信息；',
    '可中断性：用户可随时暂停或取消正在执行的任务，系统应优雅地完成当前原子操作后停止；',
    '跨会话记忆：系统记住用户偏好和历史操作，避免重复询问相同的上下文信息。',
]:
    add_para(doc, f'  ● {ux}')

add_heading_styled(doc, '4.4 可靠性需求', 2)
for rel in [
    '系统可用性：Agent核心服务可用性 ≥ 99%（7×24小时运行，不含计划维护）；',
    '故障恢复：对于可恢复的故障（如UI元素未找到），自动恢复成功率 ≥ 70%；',
    '数据持久性：任务执行记录和审计日志持久化存储，不因进程重启丢失；',
    '优雅降级：当MCP Server或视觉模型不可用时，系统应报告明确的错误信息而非崩溃；',
    '无副作用终止：Agent进程被强制终止时，已完成的文件操作不丢失，系统环境不被破坏。',
]:
    add_para(doc, f'  ● {rel}')

add_heading_styled(doc, '4.5 可维护性需求', 2)
for maint in [
    '模块化架构：各功能模块（技能库、Agent运行时、感知层、控制层）松耦合，通过接口通信，支持独立开发和测试；',
    '技能可扩展：通过YAML声明文件新增技能，无需修改Agent核心代码；',
    '配置外部化：所有配置项（模型参数、MCP Server地址、数据库路径等）通过YAML配置文件管理，不硬编码；',
    'Docker化部署：核心服务（Ollama、Jaeger、ChromaDB）通过Docker Compose一键编排启动；',
    '日志可读：结构化日志同时输出到控制台（开发环境）和文件（生产环境），日志轮转保留最近7天记录。',
]:
    add_para(doc, f'  ● {maint}')

doc.add_page_break()

# ============================================================
# CHAPTER 5: DATA REQUIREMENTS
# ============================================================
add_heading_styled(doc, '第五章 数据需求', 1)

add_heading_styled(doc, '5.1 数据实体关系', 2)
add_para(doc, '系统涉及以下核心数据实体：', indent=True)
add_table(doc,
    ['实体', '存储方式', '关键字段', '说明'],
    [
        ['SkillNode', 'YAML文件 + 内存注册表', 'id, name, description, embedding, parameters, executable_ref', '技能定义'],
        ['SkillEmbedding', 'ChromaDB', 'skill_id, embedding_vector, metadata', '技能语义向量'],
        ['Task', 'SQLite', 'task_id, user_input, status, created_at, completed_at', '任务记录'],
        ['TaskStep', 'SQLite', 'step_id, task_id, skill_id, params, status, started_at, ended_at', '任务步骤'],
        ['ExecutionTrace', 'SQLite + Jaeger', 'span_id, parent_span_id, skill_id, start_time, end_time, status', '执行追踪'],
        ['Screenshot', '文件系统', 'task_id, step_id, type(pre/post), file_path', '截图文件'],
        ['AuditLog', '文件系统 (JSONL)', 'timestamp, operation, target, result, trace_id', '审计日志'],
    ]
)

add_heading_styled(doc, '5.2 数据库选型', 2)
add_table(doc,
    ['数据库', '用途', '选型理由'],
    [
        ['ChromaDB', '技能语义向量存储与检索', '轻量嵌入式向量数据库，Python原生支持，零配置，适合小规模技能库'],
        ['SQLite', '任务记录、步骤记录、执行追踪', '零配置嵌入式关系数据库，适合单机Agent场景，数据量小（< 1GB）'],
        ['文件系统 (JSONL)', '审计日志、应用配置', '追加写入极快、人类可读、易于解析和分析'],
    ]
)

doc.add_page_break()

# ============================================================
# CHAPTER 6: EXTERNAL INTERFACES
# ============================================================
add_heading_styled(doc, '第六章 外部接口需求', 1)

add_heading_styled(doc, '6.1 用户接口', 2)
add_para(doc, '系统提供以下用户接口：', indent=True)
add_table(doc,
    ['接口类型', '说明', '使用者'],
    [
        ['Web控制台', '基于Flask/FastAPI的Web界面，提供任务提交、执行监控、历史查询功能', '终端用户'],
        ['CLI命令行', '提供命令行接口，支持快速提交任务和查看结果', '开发者、高级用户'],
        ['REST API', '提供标准RESTful API接口，支持第三方系统集成', '外部系统'],
    ]
)

add_heading_styled(doc, '6.2 软件接口', 2)
add_table(doc,
    ['接口', '协议', '方向', '说明'],
    [
        ['ZPIT-desktop-MCP', 'MCP (stdio)', '调用', '桌面控制服务：截图、键鼠、窗口管理'],
        ['Ollama', 'REST API (localhost:11434)', '调用', '本地视觉模型推理服务'],
        ['GPT-4o API', 'HTTPS REST API', '调用', '云端视觉/语言模型推理'],
        ['ChromaDB', '嵌入式SDK', '内部', '技能向量存储与检索'],
        ['SQLite', '嵌入式SQL', '内部', '任务数据持久化'],
        ['Jaeger', 'OTLP (gRPC)', '上报', '链路追踪数据导出'],
    ]
)

add_heading_styled(doc, '6.3 通信协议', 2)
add_table(doc,
    ['协议', '用途', '说明'],
    [
        ['MCP/stdio', 'Agent ↔ MCP Server', '标准输入输出流JSON-RPC通信，进程内延迟极低'],
        ['OTLP/gRPC', 'Agent → Jaeger', 'OpenTelemetry标准导出协议，支持批量压缩上报'],
        ['HTTP/HTTPS', 'Agent → 视觉模型API', '标准REST API调用，JSON格式请求/响应'],
        ['HTTP/WebSocket', 'Browser ↔ Web控制台', 'Web界面实时推送执行状态'],
    ]
)

doc.add_page_break()

# ============================================================
# CHAPTER 7: USE CASE SCENARIOS
# ============================================================
add_heading_styled(doc, '第七章 用例场景', 1)

add_heading_styled(doc, '7.1 场景一：桌面文件整理', 2)
add_para(doc, "用例描述：用户输入\"将桌面上最近3天创建的文件移动到新建的‘临时文件’文件夹\"，Agent自主完成桌面文件整理。", indent=True)
add_para(doc, '前置条件：', indent=True)
for pre in ['桌面文件夹路径可访问；', '用户对桌面文件夹有读写权限；', 'ZPIT-desktop-MCP服务正常运行；', '技能库已加载文件管理类技能。']:
    add_para(doc, f'  ● {pre}')
add_para(doc, '执行流程：', indent=True)
add_table(doc,
    ['步骤', '技能', '操作', '预期结果'],
    [
        ['1', 'open_file_explorer', '打开资源管理器，定位到桌面路径', '资源管理器窗口打开，显示桌面文件'],
        ['2', 'list_files', '列出桌面所有文件及其创建时间', '获取文件列表JSON，含文件名、路径、创建时间'],
        ['3', '（Agent内部筛选）', '根据当前日期计算3天前时间戳，筛选符合条件的文件', '符合条件的文件列表'],
        ['4', 'create_folder', '在桌面创建"临时文件"文件夹', '文件夹创建成功'],
        ['5', 'move_file (循环)', '遍历符合条件的文件，逐一移动到"临时文件"文件夹', '所有文件移动成功'],
        ['6', '（验证）', '检查目标文件夹内容，验证文件数量一致', '移动前后文件数量匹配'],
    ]
)
add_para(doc, '后置条件：符合条件的桌面文件全部移动到"临时文件"文件夹，原位置文件已不存在。', indent=True)

add_heading_styled(doc, '7.2 场景二：浏览器数据采集', 2)
add_para(doc, "用例描述：用户输入\"打开百度，搜索'2026年AI Agent发展趋势'，保存第一页搜索结果链接到文本文件\"，Agent自主完成浏览器操作。", indent=True)
add_para(doc, '前置条件：', indent=True)
for pre in ['Chrome/Edge浏览器已安装；', '系统已连接互联网；', 'ZPIT-desktop-MCP服务正常运行；', '技能库已加载浏览器自动化类技能。']:
    add_para(doc, f'  ● {pre}')
add_para(doc, '执行流程：', indent=True)
add_table(doc,
    ['步骤', '技能', '操作', '预期结果'],
    [
        ['1', 'open_browser', '打开Chrome浏览器', '浏览器窗口打开'],
        ['2', 'navigate_to', '导航到 www.baidu.com', '百度首页加载完成'],
        ['3', 'fill_form', '在搜索框中输入"2026年AI Agent发展趋势"', '搜索框显示输入文本'],
        ['4', 'click_element', '点击"百度一下"按钮', '搜索结果页加载完成'],
        ['5', 'scroll_page', '向下滚动页面获取更多结果', '更多搜索结果可见'],
        ['6', 'get_page_content', '提取搜索结果页中各条结果的标题和链接', '结构化数据（标题-链接对）'],
        ['7', '（保存）', '将提取的链接列表写入桌面文本文件', '文件成功创建并包含搜索链接'],
    ]
)

doc.add_page_break()

add_heading_styled(doc, '7.3 场景三：Word文档编辑', 2)
add_para(doc, "用例描述：用户输入\"打开桌面上的'报告.docx'，将标题改为'2026年度总结'并保存\"，Agent自主完成文档编辑。", indent=True)
add_para(doc, '前置条件与执行流程（略，格式同场景一/二）。', indent=True)

doc.add_page_break()

# ============================================================
# CHAPTER 8: ACCEPTANCE CRITERIA
# ============================================================
add_heading_styled(doc, '第八章 验收标准', 1)
add_para(doc, '本系统的验收标准分为功能验收和性能验收两部分。', indent=True)

add_heading_styled(doc, '8.1 功能验收清单', 2)
add_table(doc,
    ['编号', '验收项', '验收标准', '验收方法'],
    [
        ['AC-01', '技能定义与加载', '16个预置技能从YAML正确加载，解析无错误', '运行skill_loader测试脚本'],
        ['AC-02', '技能向量检索', '输入自然语言任务描述，返回Top-5相关技能，相似度排序合理', '人工评估10个不同任务查询'],
        ['AC-03', '任务规划', '输入自然语言任务，LLM输出结构化的步骤序列JSON', '检查输出格式正确性'],
        ['AC-04', '参数实例化', '对缺失参数能从UI上下文推断或提示用户补充', '测试3个含缺失参数的场景'],
        ['AC-05', '技能执行', '单步技能（打开文件管理器、创建文件夹等）执行成功', '逐一执行16个预置技能'],
        ['AC-06', '多步骤任务', '端到端完成"桌面文件整理"场景（≥5步）', '人工观察执行过程'],
        ['AC-07', '故障恢复（L1）', '模拟元素未找到，验证自动重试3次', '注入UI元素不可见故障'],
        ['AC-08', '故障恢复（L2）', '模拟技能超时，验证切换到替代技能', '注入技能执行超时故障'],
        ['AC-09', 'Web控制台', '通过Web界面提交任务、查看执行过程、查询历史', '手动操作Web界面'],
        ['AC-10', 'OpenTelemetry追踪', 'Jaeger UI中可查看完整调用链', '执行任务后查看Jaeger'],
        ['AC-11', 'WindowsAgentArena', '运行基准测试并输出成功率等指标', '运行基准测试脚本'],
        ['AC-12', '审计日志', '所有桌面操作有结构化日志记录', '检查日志文件内容'],
    ]
)

add_heading_styled(doc, '8.2 性能验收标准', 2)
add_table(doc,
    ['编号', '指标', '合格标准', '测试方法'],
    [
        ['PA-01', '任务成功率', '≥ 55% (WindowsAgentArena)', '运行标准测试集'],
        ['PA-02', '平均任务完成时间', '< 3分钟/任务', '计时统计'],
        ['PA-03', '故障恢复率', '> 70%', '统计故障恢复成功/失败次数'],
        ['PA-04', '技能检索延迟', '< 100ms', '性能测试脚本'],
        ['PA-05', '截图延迟', '< 200ms', '性能测试脚本'],
    ]
)

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_heading_styled(doc, '附录A：技能YAML定义示例', 1)
add_para(doc, '以下为一个完整的技能YAML定义示例：', indent=True)
code_text = '''id: open_file_explorer
name: "打开文件管理器"
description: "打开Windows资源管理器并导航到指定路径。如果未指定路径，则打开默认位置（此电脑）。"
category: file_ops
parameters:
  - name: path
    type: string
    required: false
    default: "::{20D04FE0-3AEA-1069-A2D8-08002B30309D}"
    description: "资源管理器初始路径，默认为'此电脑'"
preconditions:
  - "操作系统为Windows 10/11"
postconditions:
  - "资源管理器窗口已打开并处于前台"
  - "窗口标题栏显示指定的文件夹名称"
executable_ref: "skills.file_ops.open_file_explorer.execute"
timeout_ms: 10000
retry_policy:
  max_retries: 2
  backoff_factor: 1.5
  retry_on:
    - "WindowNotFoundError"
    - "TimeoutError"'''
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.0
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

add_heading_styled(doc, '附录B：任务规划输出JSON示例', 1)
code_text2 = '''{
  "task_id": "task_20260629_001",
  "user_input": "将桌面上最近3天创建的文件移动到新建的临时文件夹",
  "created_at": "2026-06-29T10:30:00Z",
  "estimated_duration_s": 45,
  "steps": [
    {
      "step_id": 1,
      "skill_id": "open_file_explorer",
      "params": {"path": "C:\\\\Users\\\\user\\\\Desktop"},
      "expected_precondition": "explorer_window_ready",
      "expected_postcondition": "desktop_directory_visible"
    },
    {
      "step_id": 2,
      "skill_id": "list_files",
      "params": {"path": "C:\\\\Users\\\\user\\\\Desktop", "pattern": "*.*"},
      "expected_postcondition": "file_list_obtained"
    },
    {
      "step_id": 3,
      "skill_id": "create_folder",
      "params": {"path": "C:\\\\Users\\\\user\\\\Desktop", "name": "临时文件"},
      "expected_postcondition": "folder_created"
    },
    {
      "step_id": 4,
      "skill_id": "move_file",
      "params": {
        "source": "${step2.output.filtered_files}",
        "dest": "C:\\\\Users\\\\user\\\\Desktop\\\\临时文件"
      },
      "loop_over": "step2.output.filtered_files",
      "expected_postcondition": "all_files_moved"
    }
  ]
}'''
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.0
run = p.add_run(code_text2)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

add_heading_styled(doc, '附录C：OpenTelemetry Span属性定义', 1)
add_table(doc,
    ['Span名称', '层级', '关键属性'],
    [
        ['task.execute', 'Root', 'task_id, user_input, estimated_steps'],
        ['step.execute', 'Child of task', 'step_id, skill_id, step_index, total_steps'],
        ['skill.retrieve', 'Child of step', 'query_text, top_k, candidate_count'],
        ['skill.execute', 'Child of step', 'skill_id, skill_name, params_json'],
        ['action.screenshot', 'Child of skill', 'type(pre/post), region, file_path'],
        ['action.uia_query', 'Child of skill', 'query_condition, result_element_count'],
        ['action.mcp_call', 'Child of skill', 'tool_name, tool_params, mcp_server'],
        ['action.vision_infer', 'Child of skill', 'model_name, image_size, prompt'],
        ['recovery.retry', 'Child of skill', 'attempt_number, max_attempts, error_type'],
        ['recovery.fallback', 'Child of skill', 'original_skill_id, fallback_skill_id'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = 'docs/需求规格说明书_CUA-Skill-Agent_V1.0.docx'
doc.save(output_path)
print(f'[OK] SRS saved to: {output_path}')
print(f'Estimated pages: ~25+ pages')
