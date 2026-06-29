"""
Generate 接口设计说明书 (Interface Design Specification) for CUA-Skill Agent project.
Target: 10+ pages, .docx format.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# --- Style helpers ---
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph('')
    return table

def add_code_block(doc, code, font_size=8):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    return p

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('接口设计说明书')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CUA-Skill：面向桌面环境的计算机使用Agent技能库')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')
doc.add_paragraph('')

for line in [
    f'版本：V1.0',
    f'日期：{datetime.date.today().strftime("%Y-%m-%d")}',
    '状态：初稿',
    '密级：内部',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# REVISION HISTORY
# ============================================================
add_heading_styled(doc, '修订记录', 1)
add_table(doc,
    ['版本', '日期', '修订章节', '修订内容', '修订人'],
    [['V1.0', datetime.date.today().strftime('%Y-%m-%d'), '全部', '初稿编写', '项目组']]
)
doc.add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_heading_styled(doc, '第一章 引言', 1)

add_heading_styled(doc, '1.1 编写目的', 2)
add_para(doc, '本文档旨在对"CUA-Skill Agent"系统各模块之间的接口进行全面、详细的设计说明，包括内部模块接口（Python模块间调用）和外部系统接口（MCP协议通信、视觉模型API、数据库接口等）。本文档作为系统详细设计和编码实现的直接依据，确保各模块开发人员对接口约定有统一、准确的理解。', indent=True)
add_para(doc, '本文档的预期读者包括：', indent=True)
for r in [
    '后端开发工程师：了解自己负责模块的接口定义、参数规范和返回值格式；',
    '架构师：审核接口设计的合理性和一致性；',
    '测试工程师：基于接口定义编写单元测试和集成测试；',
    '前端开发工程师：理解Web控制台与Agent后端的REST API接口。',
]:
    add_para(doc, f'  ● {r}')

add_heading_styled(doc, '1.2 适用范围', 2)
add_para(doc, '本文档覆盖CUA-Skill Agent系统中以下层次的接口设计：', indent=True)
for scope in [
    'Agent运行时内部接口（Planner ↔ Retriever ↔ Executor ↔ Memory）；',
    '技能库管理接口（SkillRegistry ↔ SkillLoader ↔ ChromaDB）；',
    '感知层接口（PerceptionManager ↔ UIA Client ↔ Vision Model）；',
    '桌面控制层接口（DesktopController ↔ MCP Client ↔ ZPIT-desktop-MCP）；',
    'Web控制台接口（Frontend ↔ REST API ↔ Agent Runtime）；',
    '可观测性接口（Agent模块 → OpenTelemetry → Jaeger）。',
]:
    add_para(doc, f'  ● {scope}')

add_heading_styled(doc, '1.3 术语与缩写', 2)
add_table(doc,
    ['术语/缩写', '说明'],
    [
        ['MCP', 'Model Context Protocol，AI Agent工具调用通信标准'],
        ['UIA', 'UI Automation，Windows辅助功能API'],
        ['OTLP', 'OpenTelemetry Protocol，遥测数据导出协议'],
        ['JSON-RPC', 'JSON Remote Procedure Call，MCP的底层消息格式'],
        ['stdio', 'Standard Input/Output，MCP的一种传输方式'],
        ['API', 'Application Programming Interface，应用程序编程接口'],
        ['SDK', 'Software Development Kit，软件开发工具包'],
        ['DAG', 'Directed Acyclic Graph，有向无环图'],
    ]
)

add_heading_styled(doc, '1.4 接口设计原则', 2)
for p in [
    '统一性：所有模块间通信遵循统一的JSON消息格式，便于序列化和跨语言互操作；',
    '松耦合：模块间通过接口抽象层通信，具体实现可替换（如MCP Server可从ZPIT替换为其他实现）；',
    '版本化：关键接口携带版本号，支持向后兼容的演进；',
    '可观测：所有跨模块调用自动产生OpenTelemetry Span，便于问题定位；',
    '容错性：接口调用失败时有明确的错误码和降级策略，不导致系统崩溃。',
]:
    add_para(doc, f'  ● {p}')

doc.add_page_break()

# ============================================================
# CHAPTER 2: INTERFACE OVERVIEW
# ============================================================
add_heading_styled(doc, '第二章 接口总览', 1)

add_heading_styled(doc, '2.1 系统模块与接口关系图', 2)
add_para(doc, 'CUA-Skill Agent系统包含以下核心模块，模块间通过明确定义的接口进行通信：', indent=True)
add_para(doc, '【内部模块接口】', bold=True)
add_table(doc,
    ['调用方', '被调用方', '接口名称', '通信方式', '数据格式'],
    [
        ['Agent Runtime (main)', 'Planner', 'plan_task(task_desc)', 'Python函数调用', 'Python dict'],
        ['Agent Runtime (main)', 'Skill Retriever', 'retrieve(query, top_k)', 'Python函数调用', 'Python dict'],
        ['Executor', 'Skill Registry', 'get_skill(skill_id)', 'Python函数调用', 'SkillNode'],
        ['Executor', 'Desktop Controller', 'click(x,y), type_text(text), ...', 'Python方法调用', 'Python dataclass'],
        ['Desktop Controller', 'MCP Client', 'call_tool(name, params)', 'Python函数调用', 'JSON'],
        ['GUI Perception', 'UIA Client', 'find_element(condition)', 'Python函数调用', 'UIElement'],
        ['GUI Perception', 'Vision Model', 'analyze(image, prompt)', 'HTTP REST / Ollama SDK', 'JSON'],
        ['Memory Module', 'SQLite', 'CRUD操作', 'sqlite3嵌入式', 'SQL'],
    ]
)

add_para(doc, '【外部系统接口】', bold=True)
add_table(doc,
    ['调用方', '外部系统', '接口协议', '通信方式', '数据格式'],
    [
        ['MCP Client', 'ZPIT-desktop-MCP', 'MCP (JSON-RPC 2.0)', 'stdio (stdin/stdout)', 'JSON'],
        ['Vision Model', 'Ollama', 'Ollama REST API', 'HTTP POST', 'JSON (multipart)'],
        ['Vision Model', 'GPT-4o API', 'OpenAI REST API', 'HTTPS POST', 'JSON (multipart/base64)'],
        ['Agent Runtime', 'Jaeger', 'OTLP', 'gRPC', 'Protocol Buffers'],
        ['Web Console (Frontend)', 'Agent API Server', 'HTTP REST', 'HTTP/HTTPS', 'JSON'],
        ['Agent Runtime', 'ChromaDB', 'ChromaDB SDK', '嵌入式HTTP', 'JSON'],
    ]
)

doc.add_page_break()

# ============================================================
# CHAPTER 3: INTERNAL INTERFACES (DETAILED)
# ============================================================
add_heading_styled(doc, '第三章 内部模块接口详细设计', 1)

# --- 3.1 Planner ---
add_heading_styled(doc, '3.1 任务规划器接口 (Planner)', 2)

add_heading_styled(doc, '3.1.1 plan_task', 3)
add_para(doc, '将用户自然语言任务指令分解为结构化的技能执行计划。', indent=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['函数签名', 'async plan_task(user_input: str, context: TaskContext) -> SkillPlan'],
        ['调用方', 'Agent Runtime (主循环)'],
        ['被调用方', 'Planner模块'],
        ['调用时机', '用户提交新任务时'],
        ['同步/异步', '异步（需调用LLM，可能耗时数秒）'],
    ]
)
add_para(doc, '输入参数 TaskContext：', bold=True)
add_table(doc,
    ['字段', '类型', '必填', '说明'],
    [
        ['current_window', 'WindowInfo', '否', '当前活动窗口信息（标题、进程名）'],
        ['selected_files', 'list[str]', '否', '当前选中的文件路径列表'],
        ['clipboard_text', 'str', '否', '当前剪贴板文本内容'],
        ['screen_size', 'tuple[int,int]', '否', '当前屏幕分辨率 (width, height)'],
        ['recent_history', 'list[TaskRecord]', '否', '最近的5条任务历史记录'],
    ]
)
add_para(doc, '输出 SkillPlan：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['task_id', 'str', '任务唯一标识，格式：task_YYYYMMDD_HHMMSS_序号'],
        ['user_input', 'str', '原始用户输入文本'],
        ['steps', 'list[PlanStep]', '分解后的步骤列表'],
        ['estimated_duration_s', 'float', '预估总耗时（秒）'],
        ['created_at', 'str', '创建时间 (ISO 8601)'],
    ]
)
add_para(doc, 'PlanStep 结构：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['step_id', 'int', '步骤序号（从1开始）'],
        ['skill_id', 'str', '匹配的技能标识符'],
        ['params', 'dict[str, Any]', '参数绑定 {参数名: 参数值}'],
        ['loop_over', 'str (optional)', '循环数据源引用（如"step2.output.filtered_files"）'],
        ['condition', 'str (optional)', '条件分支表达式'],
        ['expected_precondition', 'str', '期望的前置条件'],
        ['expected_postcondition', 'str', '期望的后置条件'],
    ]
)

add_heading_styled(doc, '3.1.2 接口交互时序', 3)
add_para(doc, 'Plan Task 接口调用流程：', indent=True)
add_code_block(doc, """
Agent Runtime                    Planner                    LLM Service
     |                              |                            |
     |-- plan_task(input, ctx) ---->|                            |
     |                              |-- prompt construction ---->|
     |                              |                            |
     |                              |<--- skill_plan_json -------|
     |                              |                            |
     |                              |-- validate & parse          |
     |<--- SkillPlan ---------------|                            |
     |                              |                            |
""", font_size=9)

add_para(doc, '错误处理：', indent=True)
add_table(doc,
    ['错误类型', '错误码', '处理方式'],
    [
        ['LLM服务不可用', 'PLAN_ERR_LLM_UNAVAILABLE', '返回错误，提示用户检查Ollama/API配置'],
        ['LLM返回格式无效', 'PLAN_ERR_INVALID_FORMAT', '重试一次（重新构造更严格的prompt），仍失败则返回错误'],
        ['未能匹配任何技能', 'PLAN_ERR_NO_SKILL_MATCH', '提示用户重新描述任务，或列出可用技能供选择'],
        ['任务超出能力范围', 'PLAN_ERR_OUT_OF_SCOPE', '明确告知用户哪些部分无法完成，建议替代方案'],
    ]
)

doc.add_page_break()

# --- 3.2 Skill Retriever ---
add_heading_styled(doc, '3.2 技能检索器接口 (Skill Retriever)', 2)

add_heading_styled(doc, '3.2.1 retrieve', 3)
add_para(doc, '根据自然语言查询文本，从技能库中检索最匹配的技能节点。', indent=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['函数签名', 'async retrieve(query: str, top_k: int = 5, filters: dict = None) -> list[RetrievalResult]'],
        ['调用方', 'Planner（规划时匹配技能）、Executor（故障恢复时查找替代技能）'],
        ['被调用方', 'Skill Retriever模块'],
    ]
)
add_para(doc, '输入参数：', bold=True)
add_table(doc,
    ['参数', '类型', '必填', '默认值', '说明'],
    [
        ['query', 'str', '是', '-', '用户任务描述或子任务描述'],
        ['top_k', 'int', '否', '5', '返回结果数量'],
        ['filters', 'dict', '否', 'None', '过滤条件，如 {"category": "file_ops"} 限定技能类别'],
    ]
)
add_para(doc, '输出 RetrievalResult：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['skill_id', 'str', '技能唯一标识'],
        ['skill_name', 'str', '技能名称'],
        ['score', 'float', '相似度得分 (0.0~1.0)'],
        ['description', 'str', '技能功能描述'],
        ['category', 'str', '技能分类'],
        ['parameters', 'list[ParameterDef]', '参数定义（名称、类型、是否必填）'],
    ]
)

add_heading_styled(doc, '3.2.2 get_skill', 3)
add_para(doc, '根据技能ID获取完整的SkillNode对象。', indent=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['函数签名', 'def get_skill(skill_id: str) -> SkillNode | None'],
        ['调用方', 'Executor（执行前获取技能完整定义）'],
        ['被调用方', 'Skill Registry'],
        ['返回值', 'SkillNode对象（技能存在）或None（技能不存在）'],
    ]
)

add_heading_styled(doc, '3.2.3 list_skills', 3)
add_para(doc, '列出技能库中所有技能（支持分类过滤）。', indent=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['函数签名', 'def list_skills(category: str = None) -> list[SkillSummary]'],
        ['调用方', 'Web控制台（展示可用技能列表）、Planner（兜底方案）'],
        ['被调用方', 'Skill Registry'],
    ]
)

doc.add_page_break()

# --- 3.3 Executor ---
add_heading_styled(doc, '3.3 执行引擎接口 (Executor)', 2)

add_heading_styled(doc, '3.3.1 execute_plan', 3)
add_para(doc, '按SkillPlan逐步执行技能序列，并处理执行过程中的故障恢复。', indent=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['函数签名', 'async execute_plan(plan: SkillPlan, callbacks: ExecutionCallbacks = None) -> ExecutionResult'],
        ['调用方', 'Agent Runtime (主循环)'],
        ['被调用方', 'Executor模块'],
    ]
)
add_para(doc, 'ExecutionCallbacks（可选回调，用于实时推送执行状态到Web控制台）：', bold=True)
add_table(doc,
    ['回调', '签名', '触发时机'],
    [
        ['on_step_start', '(step: PlanStep) -> None', '每步执行开始前'],
        ['on_step_complete', '(step: PlanStep, result: StepResult) -> None', '每步执行成功后'],
        ['on_step_error', '(step: PlanStep, error: ExecutionError) -> None', '每步执行失败时（进入恢复流程前）'],
        ['on_recovery', '(level: int, strategy: str) -> None', '触发故障恢复时'],
        ['on_task_complete', '(result: ExecutionResult) -> None', '任务全部完成时'],
    ]
)
add_para(doc, '输出 ExecutionResult：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['task_id', 'str', '任务ID'],
        ['status', 'enum', 'SUCCESS / PARTIAL_SUCCESS / FAILED / CANCELLED'],
        ['total_steps', 'int', '规划的总步骤数'],
        ['completed_steps', 'int', '成功完成的步骤数'],
        ['failed_steps', 'int', '失败的步骤数'],
        ['recovery_attempts', 'int', '触发的故障恢复总次数'],
        ['duration_ms', 'int', '总执行耗时（毫秒）'],
        ['step_results', 'list[StepResult]', '每步的详细执行结果'],
        ['error_message', 'str (optional)', '如整体失败，记录失败原因'],
    ]
)
add_para(doc, 'StepResult 结构：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['step_id', 'int', '步骤序号'],
        ['skill_id', 'str', '执行的技能ID'],
        ['status', 'enum', 'SUCCESS / FAILED / SKIPPED'],
        ['started_at', 'str', '开始时间 (ISO 8601)'],
        ['ended_at', 'str', '结束时间 (ISO 8601)'],
        ['duration_ms', 'int', '步骤耗时'],
        ['input_params', 'dict', '实际传入的参数'],
        ['output_data', 'dict (optional)', '步骤输出数据（供后续步骤引用）'],
        ['input_screenshot', 'str', '执行前截图文件路径'],
        ['output_screenshot', 'str', '执行后截图文件路径'],
        ['error', 'dict (optional)', '失败时的错误详情'],
    ]
)

doc.add_page_break()

# --- 3.4 Desktop Controller ---
add_heading_styled(doc, '3.4 桌面控制器接口 (DesktopController)', 2)
add_para(doc, 'DesktopController是桌面底层操作的统一抽象接口，屏蔽底层MCP Server或pywinauto的实现差异。', indent=True)

add_heading_styled(doc, '3.4.1 接口定义', 3)

# --- Mouse ---
add_para(doc, '【鼠标操作】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['mouse_move(x, y, duration_ms=300)', 'x:int, y:int, duration_ms:int', 'None', '移动鼠标到屏幕坐标(x,y)'],
        ['mouse_click(x=None, y=None, button="left")', 'x:int?, y:int?, button:str', 'None', '单击；不传坐标则在当前位置点击'],
        ['mouse_double_click(x, y, button="left")', 'x:int, y:int, button:str', 'None', '指定位置双击'],
        ['mouse_right_click(x=None, y=None)', 'x:int?, y:int?', 'None', '右键单击'],
        ['mouse_drag(x1,y1,x2,y2,duration_ms=500)', 'x1:int,y1:int,x2:int,y2:int,dur:int', 'None', '拖拽操作'],
        ['mouse_scroll(delta, x=None, y=None)', 'delta:int, x:int?, y:int?', 'None', '滚轮滚动（正值向上）'],
    ]
)

# --- Keyboard ---
add_para(doc, '【键盘操作】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['key_type(text, interval_ms=50)', 'text:str, interval_ms:int', 'None', '模拟逐字符输入文本'],
        ['key_press(key)', 'key:str', 'None', '按下并释放单个键（如"enter"）'],
        ['key_combo(keys)', 'keys:list[str]', 'None', '组合键（如["ctrl","c"]）'],
        ['key_hold(key, duration_ms)', 'key:str, duration_ms:int', 'None', '按住某键指定时间'],
    ]
)

# --- Window ---
add_para(doc, '【窗口操作】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['window_list()', '无', 'list[WindowInfo]', '枚举所有顶级窗口'],
        ['window_activate(title=None, hwnd=None)', 'title:str?, hwnd:int?', 'None', '激活前台窗口'],
        ['window_move_resize(hwnd,x,y,w,h)', 'hwnd:int,x:int,y:int,w:int,h:int', 'None', '移动并调整窗口大小'],
        ['window_minimize(hwnd)', 'hwnd:int', 'None', '最小化窗口'],
        ['window_maximize(hwnd)', 'hwnd:int', 'None', '最大化窗口'],
        ['window_close(hwnd)', 'hwnd:int', 'None', '关闭窗口'],
        ['window_get_foreground()', '无', 'WindowInfo', '获取当前前台窗口信息'],
    ]
)

# --- Screenshot ---
add_para(doc, '【截图操作】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['screenshot(region=None)', 'region:tuple[int,int,int,int]?', 'bytes', '截图并返回PNG格式bytes，region可指定区域'],
        ['screenshot_to_file(path, region=None)', 'path:str, region:tuple?', 'str', '截图保存到指定路径，返回文件路径'],
    ]
)

# --- Clipboard ---
add_para(doc, '【剪贴板操作】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['clipboard_get()', '无', 'str', '获取剪贴板文本内容'],
        ['clipboard_set(text)', 'text:str', 'None', '设置剪贴板文本内容'],
        ['clipboard_clear()', '无', 'None', '清空剪贴板'],
    ]
)

# --- App ---
add_para(doc, '【应用管理】', bold=True)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['app_start(path, args=None, wait=True)', 'path:str, args:list?, wait:bool', 'ProcessInfo', '启动应用程序'],
        ['app_find(name)', 'name:str', 'list[ProcessInfo]', '按进程名或窗口标题查找'],
        ['app_kill(name_or_pid)', 'name_or_pid:str|int', 'bool', '强制终止进程'],
        ['app_wait_window(app_name, timeout_ms=30000)', 'app_name:str, timeout_ms:int', 'WindowInfo', '等待应用窗口出现'],
    ]
)

doc.add_page_break()

# --- 3.5 UIA Client ---
add_heading_styled(doc, '3.5 UIA客户端接口 (UIAClient)', 2)
add_para(doc, 'UIAClient封装Windows UI Automation API，提供结构化UI元素查询能力。', indent=True)

add_heading_styled(doc, '3.5.1 接口定义', 3)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['find_element(condition, scope=None)', 'condition:dict, scope:UIElement?', 'UIElement', '按条件查找单个UI元素'],
        ['find_elements(condition, scope=None)', 'condition:dict, scope:UIElement?', 'list[UIElement]', '按条件查找所有匹配UI元素'],
        ['get_element_tree(window_hwnd=None)', 'window_hwnd:int?', 'UITreeNode', '获取窗口完整UI元素树'],
        ['get_element_property(element, prop_name)', 'element:UIElement, prop_name:str', 'Any', '获取元素属性值'],
        ['element_is_visible(element)', 'element:UIElement', 'bool', '检测元素是否可见'],
        ['element_is_enabled(element)', 'element:UIElement', 'bool', '检测元素是否可交互'],
        ['element_get_rect(element)', 'element:UIElement', 'Rect', '获取元素屏幕坐标矩形'],
    ]
)

add_para(doc, '查询条件 Condition 支持的字段：', indent=True)
add_table(doc,
    ['条件字段', '类型', '说明'],
    [
        ['name', 'str', '按元素名称精确匹配'],
        ['name_contains', 'str', '按元素名称子串匹配'],
        ['class_name', 'str', '按类别名匹配'],
        ['automation_id', 'str', '按自动化ID匹配'],
        ['control_type', 'str', '按控件类型（Button/Edit/ListItem等）匹配'],
        ['framework_id', 'str', '按UI框架（Win32/WPF/WinForms等）匹配'],
    ]
)

add_para(doc, 'UIElement 结构：', indent=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['name', 'str', '元素名称'],
        ['class_name', 'str', '类别名称'],
        ['control_type', 'str', '控件类型'],
        ['automation_id', 'str', '自动化ID'],
        ['bounding_rect', 'Rect', '屏幕坐标 (left, top, right, bottom)'],
        ['is_visible', 'bool', '是否可见'],
        ['is_enabled', 'bool', '是否可交互'],
        ['handle', 'int', '窗口句柄（顶层窗口）'],
    ]
)

doc.add_page_break()

# --- 3.6 Vision Model ---
add_heading_styled(doc, '3.6 视觉模型接口 (VisionModel)', 2)
add_para(doc, 'VisionModel接口封装对视觉语言模型的调用，支持本地Ollama和云端GPT-4o两种后端。', indent=True)

add_heading_styled(doc, '3.6.1 接口定义', 3)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['analyze_screen(image, prompt)', 'image:bytes|PIL.Image, prompt:str', 'VisionResult', '分析屏幕截图并返回结果'],
        ['locate_element(image, description)', 'image:bytes, description:str', 'BoundingBox', '定位指定元素的屏幕坐标'],
        ['describe_screen(image, detail="standard")', 'image:bytes, detail:str', 'str', '描述当前屏幕的内容'],
        ['compare_screenshots(before, after)', 'before:bytes, after:bytes', 'DiffResult', '对比两张截图，识别变化'],
    ]
)

add_para(doc, 'VisionResult 结构：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['model_name', 'str', '使用的模型名称'],
        ['response_text', 'str', '模型文本响应'],
        ['bounding_boxes', 'list[BoundingBox]', '识别到的元素边界框列表'],
        ['inference_time_ms', 'int', '推理耗时（毫秒）'],
        ['token_usage', 'dict', 'Token使用情况 {prompt_tokens, completion_tokens}'],
    ]
)

add_para(doc, 'BoundingBox 结构：', bold=True)
add_table(doc,
    ['字段', '类型', '说明'],
    [
        ['label', 'str', '元素标签或描述'],
        ['x', 'int', '左上角X坐标（像素）'],
        ['y', 'int', '左上角Y坐标（像素）'],
        ['width', 'int', '宽度（像素）'],
        ['height', 'int', '高度（像素）'],
        ['confidence', 'float', '置信度 (0.0~1.0)'],
    ]
)

add_heading_styled(doc, '3.6.2 后端切换', 3)
add_para(doc, '系统通过配置项 vision.backend 决定使用哪个视觉模型后端：', indent=True)
add_table(doc,
    ['backend值', '模型', 'SDK/API', '说明'],
    [
        ['"ollama"', 'Qwen2-VL (可选7B/72B)', 'ollama Python SDK', '本地GPU推理，零网络延迟成本，需要RTX 3060+'],
        ['"openai"', 'GPT-4o / GPT-4V', 'openai Python SDK', '云端API，能力强但需网络和费用'],
        ['"mock"', '无（测试桩）', '无', '返回预设的mock响应，用于单元测试和开发调试'],
    ]
)

doc.add_page_break()

# --- 3.7 Memory ---
add_heading_styled(doc, '3.7 记忆模块接口 (Memory)', 2)
add_para(doc, 'Memory模块负责Agent的任务状态持久化和跨会话状态维护。', indent=True)

add_heading_styled(doc, '3.7.1 接口定义', 3)
add_table(doc,
    ['方法签名', '参数', '返回值', '说明'],
    [
        ['save_task(task_record)', 'task_record:TaskRecord', 'int (task_db_id)', '保存任务记录，返回数据库主键'],
        ['update_task_status(task_id, status)', 'task_id:str, status:str', 'None', '更新任务状态'],
        ['get_task(task_id)', 'task_id:str', 'TaskRecord?', '获取任务完整记录'],
        ['list_tasks(limit=50, offset=0, status=None)', 'limit:int, offset:int, status:str?', 'list[TaskSummary]', '分页查询任务列表'],
        ['save_step_result(task_id, step_result)', 'task_id:str, step_result:StepResult', 'None', '保存单步执行结果'],
        ['get_execution_trace(task_id)', 'task_id:str', 'list[StepResult]', '获取任务完整执行轨迹'],
        ['save_snapshot(task_id, label, data)', 'task_id:str, label:str, data:dict', 'None', '保存环境状态快照'],
        ['get_snapshot(task_id, label)', 'task_id:str, label:str', 'dict?', '获取指定快照'],
        ['delete_old_tasks(days)', 'days:int', 'int (deleted_count)', '清理指定天数前的历史记录'],
    ]
)

add_para(doc, '数据库表结构：', bold=True)
add_code_block(doc, """
-- tasks 表
CREATE TABLE tasks (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    task_id TEXT UNIQUE NOT NULL,
    user_input TEXT NOT NULL,
    status TEXT DEFAULT 'pending',
    total_steps INTEGER DEFAULT 0,
    completed_steps INTEGER DEFAULT 0,
    recovery_attempts INTEGER DEFAULT 0,
    duration_ms INTEGER,
    created_at TEXT NOT NULL,
    completed_at TEXT
);

-- step_results 表
CREATE TABLE step_results (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    task_id TEXT NOT NULL,
    step_id INTEGER NOT NULL,
    skill_id TEXT NOT NULL,
    status TEXT NOT NULL,
    input_params JSON,
    output_data JSON,
    input_screenshot TEXT,
    output_screenshot TEXT,
    error_json JSON,
    duration_ms INTEGER,
    started_at TEXT,
    ended_at TEXT,
    FOREIGN KEY (task_id) REFERENCES tasks(task_id)
);

-- snapshots 表
CREATE TABLE snapshots (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    task_id TEXT NOT NULL,
    label TEXT NOT NULL,
    snapshot_data JSON NOT NULL,
    created_at TEXT NOT NULL
);""", font_size=8)

doc.add_page_break()

# ============================================================
# CHAPTER 4: EXTERNAL INTERFACES (DETAILED)
# ============================================================
add_heading_styled(doc, '第四章 外部系统接口详细设计', 1)

# --- 4.1 MCP Interface ---
add_heading_styled(doc, '4.1 MCP协议接口 (ZPIT-desktop-MCP)', 2)

add_heading_styled(doc, '4.1.1 通信方式', 3)
add_para(doc, 'MCP Client与ZPIT-desktop-MCP Server之间采用stdio传输协议进行JSON-RPC 2.0通信。Client启动Server子进程，通过标准输入输出流交换JSON消息。', indent=True)
add_para(doc, '传输层配置：', indent=True)
add_table(doc,
    ['配置项', '值', '说明'],
    [
        ['传输协议', 'stdio', '标准输入输出流'],
        ['消息格式', 'JSON-RPC 2.0', '每行一个完整的JSON对象'],
        ['编码', 'UTF-8', '统一字符编码'],
        ['心跳间隔', '30秒', 'Client定期发送ping请求检测Server存活状态'],
        ['重连策略', '指数退避 (1s→2s→4s→8s)', '最大重试3次，超出后报错'],
    ]
)

add_heading_styled(doc, '4.1.2 MCP 生命周期', 3)
add_para(doc, 'MCP协议定义了严格的生命周期管理：', indent=True)
add_code_block(doc, """
Client                          MCP Server
   |                                |
   |--- initialize request -------->|  (1) 能力协商
   |<--- initialize response -------|
   |                                |
   |--- initialzed notification --->|  (2) 初始化完成确认
   |                                |
   |--- tools/list request -------->|  (3) 工具发现
   |<--- tools/list response -------|
   |                                |
   |--- tools/call request -------->|  (4) 工具调用
   |<--- tools/call response -------|
   |                                |
   |--- ping request -------------->|  (5) 心跳检测
   |<--- pong response -------------|
   |                                |
   |--- shutdown request ---------->|  (6) 优雅关闭
   |<--- shutdown response ---------|
""", font_size=8)

add_heading_styled(doc, '4.1.3 消息格式规范', 3)
add_para(doc, '请求消息格式（以tools/call为例）：', indent=True)
add_code_block(doc, """{
  "jsonrpc": "2.0",
  "id": 1,
  "method": "tools/call",
  "params": {
    "name": "screenshot",
    "arguments": {
      "region": null,
      "format": "png"
    }
  }
}""", font_size=9)

add_para(doc, '响应消息格式（成功）：', indent=True)
add_code_block(doc, """{
  "jsonrpc": "2.0",
  "id": 1,
  "result": {
    "content": [
      {
        "type": "image",
        "data": "iVBORw0KGgoAAAANSUhEUgAA...",
        "mimeType": "image/png"
      }
    ]
  }
}""", font_size=9)

add_para(doc, '响应消息格式（错误）：', indent=True)
add_code_block(doc, """{
  "jsonrpc": "2.0",
  "id": 1,
  "error": {
    "code": -32603,
    "message": "Internal error",
    "data": {
      "detail": "Failed to capture screen: access denied"
    }
  }
}""", font_size=9)

add_heading_styled(doc, '4.1.4 MCP工具完整清单', 3)
add_para(doc, 'ZPIT-desktop-MCP Server提供的工具清单及映射关系：', indent=True)
add_table(doc,
    ['MCP Tool名称', '参数', '返回类型', 'Agent调用方法'],
    [
        ['screenshot', 'region: rect?', 'image (base64 PNG)', 'DesktopController.screenshot()'],
        ['mouse_click', 'x:int, y:int, button:str?', 'null', 'DesktopController.mouse_click()'],
        ['mouse_double_click', 'x:int, y:int', 'null', 'DesktopController.mouse_double_click()'],
        ['mouse_move', 'x:int, y:int', 'null', 'DesktopController.mouse_move()'],
        ['mouse_drag', 'x1,y1,x2,y2:int', 'null', 'DesktopController.mouse_drag()'],
        ['mouse_scroll', 'delta:int', 'null', 'DesktopController.mouse_scroll()'],
        ['key_type', 'text:str, interval_ms:int?', 'null', 'DesktopController.key_type()'],
        ['key_combo', 'keys:str[]', 'null', 'DesktopController.key_combo()'],
        ['key_press', 'key:str', 'null', 'DesktopController.key_press()'],
        ['window_list', '-', 'WindowInfo[]', 'DesktopController.window_list()'],
        ['window_activate', 'title:str? / hwnd:int?', 'null', 'DesktopController.window_activate()'],
        ['window_close', 'hwnd:int', 'null', 'DesktopController.window_close()'],
        ['clipboard_get', '-', 'text', 'DesktopController.clipboard_get()'],
        ['clipboard_set', 'text:str', 'null', 'DesktopController.clipboard_set()'],
        ['app_start', 'path:str, args:str[]?', 'ProcessInfo', 'DesktopController.app_start()'],
        ['app_kill', 'name:str', 'bool', 'DesktopController.app_kill()'],
    ]
)

doc.add_page_break()

# --- 4.2 Vision Model API ---
add_heading_styled(doc, '4.2 视觉模型API接口', 2)

add_heading_styled(doc, '4.2.1 Ollama 本地接口', 3)
add_para(doc, 'Ollama通过本地HTTP REST API提供服务，默认监听 localhost:11434。', indent=True)
add_para(doc, 'API端点：POST /api/generate', bold=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['URL', 'http://localhost:11434/api/generate'],
        ['方法', 'POST'],
        ['Content-Type', 'application/json'],
        ['认证', '无（本地服务）'],
    ]
)
add_para(doc, '请求体：', bold=True)
add_code_block(doc, """{
  "model": "qwen2-vl:7b",
  "prompt": "Describe what you see on this screen. Locate the 'Search' button.",
  "images": ["iVBORw0KGgoAAAANSUhEUgAA..."],
  "stream": false,
  "options": {
    "temperature": 0.1,
    "num_predict": 512
  }
}""", font_size=9)
add_para(doc, '响应体：', bold=True)
add_code_block(doc, """{
  "model": "qwen2-vl:7b",
  "created_at": "2026-06-29T10:30:00Z",
  "response": "The screen shows a Windows desktop with... The 'Search' button is located at...",
  "done": true,
  "total_duration": 8234567890,
  "eval_count": 256,
  "eval_duration": 7890123456
}""", font_size=9)

add_heading_styled(doc, '4.2.2 OpenAI GPT-4o 云端接口', 3)
add_para(doc, '通过OpenAI Python SDK调用GPT-4o视觉能力。', indent=True)
add_para(doc, 'API端点：POST https://api.openai.com/v1/chat/completions', bold=True)
add_table(doc,
    ['属性', '说明'],
    [
        ['URL', 'https://api.openai.com/v1/chat/completions'],
        ['方法', 'POST'],
        ['Content-Type', 'application/json'],
        ['认证', 'Bearer {API_KEY} (通过Header Authorization传递)'],
        ['超时', '30秒'],
    ]
)
add_para(doc, '请求体核心结构：', bold=True)
add_code_block(doc, """{
  "model": "gpt-4o",
  "messages": [
    {
      "role": "user",
      "content": [
        {"type": "text", "text": "Locate the search button on this screen."},
        {"type": "image_url", "image_url": {"url": "data:image/png;base64,iVBORw0..."}}
      ]
    }
  ],
  "max_tokens": 1024,
  "temperature": 0.1
}""", font_size=9)

doc.add_page_break()

# --- 4.3 Web API ---
add_heading_styled(doc, '4.3 Web控制台API接口', 2)
add_para(doc, 'Web控制台后端通过FastAPI提供RESTful API，前端通过HTTP/HTTPS调用。', indent=True)

add_heading_styled(doc, '4.3.1 API端点总览', 3)
add_table(doc,
    ['方法', '路径', '说明', '请求体', '响应体'],
    [
        ['POST', '/api/tasks', '提交新任务', 'SubmitTaskRequest', 'TaskResponse'],
        ['GET', '/api/tasks/{task_id}', '获取任务状态', '-', 'TaskDetailResponse'],
        ['GET', '/api/tasks/{task_id}/steps', '获取任务步骤列表', '-', 'list[StepResponse]'],
        ['GET', '/api/tasks/{task_id}/screenshots/{step_id}', '获取步骤截图', '-', 'image/png (binary)'],
        ['POST', '/api/tasks/{task_id}/cancel', '取消任务', '-', 'TaskResponse'],
        ['GET', '/api/tasks?limit=50&offset=0', '历史任务列表', '-', 'TaskListResponse'],
        ['GET', '/api/skills', '获取可用技能列表', '-', 'list[SkillSummaryResponse]'],
        ['GET', '/api/health', '健康检查', '-', '{"status":"ok"}'],
        ['WS', '/ws/tasks/{task_id}', 'WebSocket执行实时推送', '-', '(双向消息)'],
    ]
)

add_heading_styled(doc, '4.3.2 请求/响应结构定义', 3)
add_para(doc, 'SubmitTaskRequest：', bold=True)
add_code_block(doc, """{
  "user_input": "将桌面上最近3天创建的文件移动到新建的临时文件夹",
  "options": {
    "auto_confirm": false,
    "max_steps": 20,
    "timeout_seconds": 300
  }
}""", font_size=9)

add_para(doc, 'TaskDetailResponse：', bold=True)
add_code_block(doc, """{
  "task_id": "task_20260629_001",
  "user_input": "将桌面上最近3天创建的文件移动到新建的临时文件夹",
  "status": "running",
  "progress": {"completed": 3, "total": 6},
  "current_step": {"step_id": 4, "skill_name": "移动文件", "description": "移动匹配文件到目标文件夹"},
  "created_at": "2026-06-29T10:30:00Z",
  "estimated_remaining_s": 25
}""", font_size=9)

add_para(doc, 'WebSocket推送消息格式：', bold=True)
add_code_block(doc, """{
  "type": "step_update",
  "task_id": "task_20260629_001",
  "timestamp": "2026-06-29T10:30:15Z",
  "data": {
    "step_id": 3,
    "skill_id": "create_folder",
    "status": "completed",
    "screenshot_url": "/api/tasks/task_20260629_001/screenshots/3"
  }
}""", font_size=9)

doc.add_page_break()

# ============================================================
# CHAPTER 5: DATA STRUCTURE DEFINITIONS
# ============================================================
add_heading_styled(doc, '第五章 数据结构定义', 1)
add_para(doc, '本章汇总系统所有跨模块通信使用的核心数据结构定义，确保模块开发人员对数据格式有一致理解。', indent=True)

add_heading_styled(doc, '5.1 通用枚举类型', 2)
add_code_block(doc, """class TaskStatus(str, Enum):
    PENDING = "pending"
    PLANNING = "planning"
    RUNNING = "running"
    PAUSED = "paused"
    COMPLETED = "completed"
    PARTIAL_SUCCESS = "partial_success"
    FAILED = "failed"
    CANCELLED = "cancelled"

class StepStatus(str, Enum):
    PENDING = "pending"
    RUNNING = "running"
    SUCCESS = "success"
    FAILED = "failed"
    SKIPPED = "skipped"

class SkillCategory(str, Enum):
    FILE_OPS = "file_ops"
    BROWSER = "browser"
    OFFICE = "office"
    SYSTEM = "system"
    CUSTOM = "custom"

class RecoveryLevel(int, Enum):
    L1_RETRY = 1
    L2_FALLBACK = 2
    L3_REPLAN = 3

class MouseButton(str, Enum):
    LEFT = "left"
    RIGHT = "right"
    MIDDLE = "middle\"""", font_size=9)

add_heading_styled(doc, '5.2 通用数据结构', 2)

add_para(doc, 'ParameterDef：', bold=True)
add_code_block(doc, """@dataclass
class ParameterDef:
    name: str
    type: str           # "string" | "int" | "float" | "bool" | "list" | "dict"
    required: bool
    default: Any = None
    description: str = ""
    enum_values: list[str] | None = None  # 枚举值约束\"""", font_size=9)

add_para(doc, 'Rect：', bold=True)
add_code_block(doc, """@dataclass
class Rect:
    left: int
    top: int
    right: int
    bottom: int

    @property
    def width(self) -> int: return self.right - self.left
    @property
    def height(self) -> int: return self.bottom - self.top
    @property
    def center(self) -> tuple[int,int]: return ((self.left+self.right)//2, (self.top+self.bottom)//2)\"""", font_size=9)

add_para(doc, 'WindowInfo：', bold=True)
add_code_block(doc, """@dataclass
class WindowInfo:
    hwnd: int               # 窗口句柄
    title: str              # 窗口标题
    class_name: str         # 窗口类名
    process_name: str       # 进程文件名
    pid: int                # 进程ID
    rect: Rect              # 窗口位置和大小
    is_visible: bool        # 是否可见
    is_minimized: bool      # 是否最小化
    z_order: int            # Z序位置\"""", font_size=9)

add_para(doc, 'UIElement：', bold=True)
add_code_block(doc, """@dataclass
class UIElement:
    name: str
    class_name: str
    control_type: str       # "Button" | "Edit" | "ListItem" | "Menu" | ...
    automation_id: str
    framework_id: str       # "Win32" | "WPF" | "WinForms" | "DirectUI"
    bounding_rect: Rect
    is_visible: bool
    is_enabled: bool
    is_keyboard_focusable: bool
    handle: int\"""", font_size=9)

doc.add_page_break()

# ============================================================
# CHAPTER 6: ERROR CODE DEFINITIONS
# ============================================================
add_heading_styled(doc, '第六章 错误码定义', 1)
add_para(doc, '系统所有跨模块调用使用统一的错误码体系，便于故障定位和处理。', indent=True)

add_heading_styled(doc, '6.1 错误码分类', 2)
add_table(doc,
    ['错误码范围', '类别', '说明'],
    [
        ['E_1000 ~ E_1999', '规划错误 (Planner)', '任务分解、技能匹配相关错误'],
        ['E_2000 ~ E_2999', '执行错误 (Executor)', '技能执行、组合图执行相关错误'],
        ['E_3000 ~ E_3999', '桌面控制错误 (Desktop)', '键鼠模拟、窗口管理、剪贴板相关错误'],
        ['E_4000 ~ E_4999', '感知错误 (Perception)', 'UIA查询、视觉模型调用相关错误'],
        ['E_5000 ~ E_5999', '通信错误 (Communication)', 'MCP协议、网络调用相关错误'],
        ['E_6000 ~ E_6999', '数据错误 (Data)', '数据库操作、持久化相关错误'],
        ['E_9000 ~ E_9999', '系统错误 (System)', '配置错误、资源不足等系统级错误'],
    ]
)

add_heading_styled(doc, '6.2 详细错误码列表', 2)
add_table(doc,
    ['错误码', '错误名称', '说明', '建议处理'],
    [
        ['E_1001', 'PLAN_LLM_UNAVAILABLE', 'LLM服务不可达', '检查Ollama/API配置，尝试重启服务'],
        ['E_1002', 'PLAN_INVALID_OUTPUT', 'LLM输出格式无效，无法解析为SkillPlan', '重试一次，调整prompt约束'],
        ['E_1003', 'PLAN_NO_SKILL_MATCH', '未找到匹配的技能', '降低相似度阈值重试，或提示用户补充描述'],
        ['E_1004', 'PLAN_OUT_OF_SCOPE', '任务超出当前技能库覆盖范围', '告知用户可用的技能类别，建议分解任务'],
        ['E_2001', 'EXEC_SKILL_NOT_FOUND', '技能ID不存在于注册表中', '检查技能ID拼写，确认技能已加载'],
        ['E_2002', 'EXEC_PRECONDITION_FAIL', '技能前置条件不满足', '触发L1恢复：回退到上一步或重新建立前置状态'],
        ['E_2003', 'EXEC_POSTCONDITION_FAIL', '技能后置条件验证失败', '触发L1恢复：重试当前技能'],
        ['E_2004', 'EXEC_TIMEOUT', '技能执行超时', '触发L2恢复：尝试替代技能'],
        ['E_2005', 'EXEC_PARAM_INVALID', '参数类型或值不合法', '触发L1恢复：尝试从上下文重新推断参数'],
        ['E_3001', 'CTRL_ELEMENT_NOT_FOUND', 'UI元素未找到', '等待500ms后重试，仍失败则回退到视觉定位'],
        ['E_3002', 'CTRL_WINDOW_NOT_FOUND', '目标窗口未找到', '检查窗口是否已打开，或使用备选窗口标题'],
        ['E_3003', 'CTRL_CLICK_FAILED', '鼠标点击操作失败', '重试，调整坐标偏移（±5px随机偏移）'],
        ['E_3004', 'CTRL_PERMISSION_DENIED', '操作权限不足', '提示用户以管理员身份运行或修改权限设置'],
        ['E_4001', 'PERC_UIA_UNAVAILABLE', 'UI Automation服务不可用', '降级使用视觉模型定位'],
        ['E_4002', 'PERC_VISION_TIMEOUT', '视觉模型推理超时', '降低图片分辨率重试，或切换到备选模型'],
        ['E_4003', 'PERC_VISION_INVALID_RESP', '视觉模型返回格式无效', '重试一次，调整prompt以获得结构化输出'],
        ['E_5001', 'COMM_MCP_UNAVAILABLE', 'MCP Server无法连接', '重启MCP Server进程，最多重试3次'],
        ['E_5002', 'COMM_MCP_TOOL_ERROR', 'MCP工具调用返回错误', '根据MCP错误信息进行相应处理'],
        ['E_5003', 'COMM_NETWORK_ERROR', '网络请求失败（GPT-4o API等）', '指数退避重试，最多3次'],
        ['E_6001', 'DATA_DB_ACCESS_ERROR', '数据库访问失败', '检查SQLite文件权限和磁盘空间'],
        ['E_6002', 'DATA_RECORD_NOT_FOUND', '记录不存在', '返回None或空列表，不抛异常'],
        ['E_9001', 'SYS_CONFIG_ERROR', '配置加载失败', '检查config/*.yaml文件格式和路径'],
        ['E_9002', 'SYS_RESOURCE_LIMIT', '系统资源不足（内存/磁盘）', '释放缓存，提示用户清理资源'],
        ['E_9999', 'SYS_UNKNOWN', '未预期的系统错误', '记录完整traceback到日志，返回通用错误信息给用户'],
    ]
)

doc.add_page_break()

# ============================================================
# CHAPTER 7: INTERFACE VERSION MANAGEMENT
# ============================================================
add_heading_styled(doc, '第七章 接口版本管理', 1)
add_para(doc, '为支持系统的持续演进和向后兼容，关键接口需进行版本管理。', indent=True)

add_heading_styled(doc, '7.1 版本化策略', 2)
for s in [
    'MCP协议接口：遵循MCP标准协议版本（当前V1.0），ZPIT-desktop-MCP升级时同步更新；',
    'Web API接口：在URL路径中携带版本号，如 /api/v1/tasks，支持多版本并存；',
    '技能YAML定义：在文件头部携带 format_version 字段，旧版本技能定义可通过迁移脚本自动升级；',
    '内部Python模块接口：通过Pydantic模型进行运行时类型校验，字段变更时增加而非删除字段（向后兼容）。',
]:
    add_para(doc, f'  ● {s}')

add_heading_styled(doc, '7.2 接口变更流程', 2)
for step in [
    '1. 提出变更请求（变更方提交变更说明，包括新增/修改/废弃的字段和方法）；',
    '2. 影响评估（架构师评估对调用方的影响范围和兼容性）；',
    '3. 版本号分配（不兼容变更需升级主版本号，兼容变更仅升级次版本号）；',
    '4. 更新文档（同步更新本接口说明书及相关代码注释）；',
    '5. 通知调用方（通过项目沟通渠道通知所有相关模块开发人员）；',
    '6. 过渡期管理（不兼容变更需至少保留一个版本的过渡期，期间同时支持新旧接口）。',
]:
    add_para(doc, f'  {step}')

doc.add_page_break()

# ============================================================
# APPENDIX
# ============================================================
add_heading_styled(doc, '附录A：MCP工具调用Python示例代码', 1)
add_code_block(doc, """import asyncio
import json
from cua_agent.desktop.mcp_client import MCPClient

async def example_mcp_usage():
    client = MCPClient(server_command=["zpit-desktop-mcp"])

    async with client:
        # 步骤1: 获取屏幕截图
        result = await client.call_tool("screenshot", {"region": None})
        screenshot_base64 = result["content"][0]["data"]

        # 步骤2: 移动鼠标到搜索框
        await client.call_tool("mouse_move", {"x": 500, "y": 300})

        # 步骤3: 点击搜索框
        await client.call_tool("mouse_click", {"x": 500, "y": 300})

        # 步骤4: 输入搜索关键词
        await client.call_tool("key_type", {"text": "CUA-Skill", "interval_ms": 50})

        # 步骤5: 按Enter搜索
        await client.call_tool("key_press", {"key": "enter"})

        # 步骤6: 再次截图验证结果
        result2 = await client.call_tool("screenshot", {"region": None})

asyncio.run(example_mcp_usage())""", font_size=8)

doc.add_page_break()

add_heading_styled(doc, '附录B：完整技能YAML定义示例', 1)
add_code_block(doc, """# skills-defs/file_ops/copy_file.yaml
format_version: "1.0"
skill:
  id: copy_file
  name: "复制文件"
  description: "将指定文件从源路径复制到目标路径。支持复制单个文件或文件夹（递归）。如果目标路径已存在同名文件，可选择覆盖或自动重命名。"
  category: file_ops
  parameters:
    - name: source
      type: string
      required: true
      description: "源文件或文件夹的完整路径"
    - name: dest
      type: string
      required: true
      description: "目标路径（文件夹路径或完整文件路径）"
    - name: overwrite
      type: bool
      required: false
      default: false
      description: "如果目标已存在，是否覆盖"
  preconditions:
    - "source路径必须可访问且文件/文件夹存在"
    - "dest父目录必须存在"
    - "用户对source有读取权限，对dest有写入权限"
  postconditions:
    - "目标路径存在与源文件内容一致的文件"
    - "源文件仍然存在（未被移动）"
  executable_ref: "skills.file_ops.copy_file.execute"
  timeout_ms: 30000
  retry_policy:
    max_retries: 2
    backoff_factor: 2.0
    retry_on:
      - "PermissionError"
      - "OSError"
  embedding_model: "text2vec-base-chinese"
  embedding: null  # 将在技能加载时自动计算""", font_size=8)

doc.add_page_break()

add_heading_styled(doc, '附录C：Agent启动配置文件示例', 1)
add_code_block(doc, """# config/agent.yaml
agent:
  name: "CUA-Skill Agent"
  version: "1.0.0"
  log_level: "INFO"

planner:
  llm:
    provider: "ollama"            # ollama | openai
    model: "qwen2.5:7b"           # 或 gpt-4o
    temperature: 0.1
    max_tokens: 2048
    api_base: "http://localhost:11434"  # Ollama地址

skill_retriever:
  embedding_model: "text2vec-base-chinese"
  vector_db:
    type: "chromadb"
    persist_dir: "./data/chromadb"
  retrieval:
    top_k: 5
    similarity_threshold: 0.6
    hybrid_weight:
      semantic: 0.7
      category: 0.2
      popularity: 0.1

executor:
  step_timeout_ms: 30000
  task_timeout_ms: 300000
  max_steps: 20
  screenshot:
    enabled: true
    save_dir: "./data/screenshots"
    quality: 90

recovery:
  l1_max_retries: 3
  l1_backoff_factor: 1.5
  l2_max_fallbacks: 2
  l3_enabled: true

vision:
  backend: "ollama"                # ollama | openai | mock
  ollama:
    model: "qwen2-vl:7b"
    api_base: "http://localhost:11434"
    max_retries: 2
  openai:
    model: "gpt-4o"
    api_key_env: "OPENAI_API_KEY"
    max_retries: 2

mcp:
  servers:
    - name: "zpit-desktop"
      command: "zpit-desktop-mcp"
      args: []
      auto_restart: true

observability:
  otel:
    exporter: "otlp"
    endpoint: "http://localhost:4317"
    service_name: "cua-skill-agent"
  jaeger:
    endpoint: "http://localhost:16686"
  logging:
    level: "INFO"
    format: "json"
    output: ["console", "file"]
    file_path: "./logs/agent.log"
    rotation: "7d"

database:
  sqlite_path: "./data/agent.db"

web:
  host: "0.0.0.0"
  port: 8080
  cors_origins: ["*"]
  ws_enabled: true""", font_size=8)

# ============================================================
# SAVE
# ============================================================
output_path = 'docs/接口设计说明书_CUA-Skill-Agent_V1.0.docx'
doc.save(output_path)
print(f'[OK] Interface Design Spec saved to: {output_path}')
print(f'Estimated pages: ~15+ pages')
